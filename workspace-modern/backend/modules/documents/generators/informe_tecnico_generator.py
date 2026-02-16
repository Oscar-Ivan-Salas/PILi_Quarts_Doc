# -*- coding: utf-8 -*-
"""
Generador de Informe Técnico - COMPLETO
Informe técnico profesional con análisis detallado
100% fidelidad al HTML profesional aprobado
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class InformeTecnicoGenerator(BaseDocumentGenerator):
    """Generador para informes técnicos profesionales"""
    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('INFORME TÉCNICO')
        run_titulo.font.size = Pt(30)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        titulo = self.datos.get('titulo', 'Informe Técnico')
        p_subtitulo = self.doc.add_paragraph()
        p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_subtitulo = p_subtitulo.add_run(titulo)
        run_subtitulo.font.size = Pt(16)
        run_subtitulo.font.color.rgb = self.COLOR_SECUNDARIO
        run_subtitulo.font.bold = True
        
        codigo = self.datos.get('codigo', 'INF-TEC-001')
        p_codigo = self.doc.add_paragraph()
        p_codigo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_codigo = p_codigo.add_run(f'Código: {codigo}')
        run_codigo.font.size = Pt(14)
        run_codigo.font.color.rgb = self.COLOR_SECUNDARIO
        
        self.doc.add_paragraph()
    
    def _agregar_info_general(self):
        """Agrega información general en tabla"""
        table = self.doc.add_table(rows=1, cols=2)
        
        # Cliente
        cell_cliente = table.rows[0].cells[0]
        p1 = cell_cliente.paragraphs[0]
        run1 = p1.add_run('DATOS DEL CLIENTE')
        run1.font.size = Pt(14)
        run1.font.bold = True
        run1.font.color.rgb = self.COLOR_PRIMARIO
        
        cliente_data = self.datos.get('cliente', {})
        cliente = cliente_data.get('nombre', 'Cliente') if isinstance(cliente_data, dict) else str(cliente_data)
        
        cell_cliente.add_paragraph(f'Cliente: {cliente}').runs[0].font.size = Pt(12)
        
        # Informe
        cell_inf = table.rows[0].cells[1]
        p2 = cell_inf.paragraphs[0]
        run2 = p2.add_run('DATOS DEL INFORME')
        run2.font.size = Pt(14)
        run2.font.bold = True
        run2.font.color.rgb = self.COLOR_PRIMARIO
        
        fecha = self.datos.get('fecha', '01/01/2025')
        codigo = self.datos.get('codigo', 'INF-TEC-001')
        
        cell_inf.add_paragraph(f'Fecha: {fecha}').runs[0].font.size = Pt(12)
        cell_inf.add_paragraph(f'Código: {codigo}').runs[0].font.size = Pt(12)
        
        # Bordes
        for cell in [cell_cliente, cell_inf]:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_resumen_ejecutivo(self):
        """Agrega resumen ejecutivo"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('RESUMEN EJECUTIVO')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        resumen = self.datos.get('resumen_ejecutivo', 'Resumen ejecutivo del informe...')
        p_resumen = self.doc.add_paragraph(resumen)
        p_resumen.runs[0].font.size = Pt(13)
        p_resumen.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_introduccion(self):
        """Agrega introducción"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('1. INTRODUCCIÓN')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        
        introduccion = self.datos.get('introduccion', 'Introducción del informe técnico...')
        p_intro = self.doc.add_paragraph()
        run_intro = p_intro.add_run(introduccion)
        run_intro.font.size = Pt(12)
        run_intro.font.color.rgb = RGBColor(55, 65, 81)
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_analisis_tecnico(self):
        """Agrega análisis técnico"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('2. ANÁLISIS TÉCNICO')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        analisis = self.datos.get('analisis_tecnico', 'Análisis técnico detallado...')
        p_analisis = self.doc.add_paragraph()
        run_analisis = p_analisis.add_run(analisis)
        run_analisis.font.size = Pt(12)
        run_analisis.font.color.rgb = RGBColor(55, 65, 81)
        p_analisis.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_resultados(self):
        """Agrega resultados"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('3. RESULTADOS')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        resultados = self.datos.get('resultados', 'Resultados obtenidos...')
        p_resultados = self.doc.add_paragraph()
        run_resultados = p_resultados.add_run(resultados)
        run_resultados.font.size = Pt(12)
        run_resultados.font.color.rgb = RGBColor(55, 65, 81)
        p_resultados.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_conclusiones(self):
        """Agrega conclusiones"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('CONCLUSIONES')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        conclusiones = self.datos.get('conclusiones', 'Conclusiones del análisis...')
        p_conclusiones = self.doc.add_paragraph()
        run_conclusiones = p_conclusiones.add_run(conclusiones)
        run_conclusiones.font.size = Pt(12)
        run_conclusiones.font.color.rgb = RGBColor(55, 65, 81)
        p_conclusiones.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_recomendaciones(self):
        """Agrega recomendaciones"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('RECOMENDACIONES')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        recomendaciones = self.datos.get('recomendaciones', [
            'Recomendación 1',
            'Recomendación 2',
            'Recomendación 3'
        ])
        
        for rec in recomendaciones:
            p_rec = self.doc.add_paragraph(f'● {rec}')
            p_rec.runs[0].font.size = Pt(12)
            p_rec.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            p_rec.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_paragraph()
    
    def generar(self, ruta_salida):
        """Genera el documento completo"""
        # Agregar header con logo y datos de empresa
        self._agregar_header_basico()
        self._agregar_titulo()
        self._agregar_info_general()
        self._agregar_resumen_ejecutivo()
        self._agregar_introduccion()
        self._agregar_analisis_tecnico()
        self._agregar_resultados()
        self._agregar_conclusiones()
        self._agregar_recomendaciones()
        self._agregar_firma_ingenieria() # ✅ ADDED SIGNATURE BLOCK
        self._agregar_footer_basico()
        
        self.doc.save(str(ruta_salida))
        return ruta_salida


def generar_informe_tecnico(datos, ruta_salida, opciones=None):
    """Función de entrada para generar informe técnico"""
    generator = InformeTecnicoGenerator(datos, opciones)
    return generator.generar(ruta_salida)
