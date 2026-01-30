"""
Document Generator Module - Word Generator
Enterprise-grade Word document generation using python-docx
Following clean-code and python-patterns skills
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List, Any, Optional
from datetime import datetime
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


class WordGenerator:
    """
    Professional Word document generator using python-docx.
    
    Features:
    - Custom styles
    - Tables
    - Headers/footers
    - Images
    - Professional formatting
    
    Following clean-code: Single Responsibility
    """
    
    def __init__(self):
        """Initialize Word generator"""
        logger.info("Word Generator initialized")
    
    def _setup_styles(self, doc: Document):
        """
        Setup custom document styles.
        
        Following clean-code: Separation of concerns
        """
        # Title style
        if 'CustomTitle' not in doc.styles:
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(30, 64, 175)  # Blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(20)
        
        # Heading style
        if 'CustomHeading' not in doc.styles:
            heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(16)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(59, 130, 246)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    def generate_cotizacion(
        self,
        data: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> BytesIO:
        """
        Generate professional quotation Word document.
        
        Following python-patterns: Type hints and clear structure
        
        Args:
            data: Quotation data
            output_path: Optional file path to save document
        
        Returns:
            BytesIO buffer with Word document
        """
        logger.info("Generating quotation Word document")
        
        # Create document
        doc = Document()
        self._setup_styles(doc)
        
        # Header
        self._add_header(doc, data)
        
        # Client info
        self._add_client_info(doc, data.get('cliente', {}))
        
        # Items table
        self._add_items_table(doc, data.get('items', []))
        
        # Totals
        self._add_totals(doc, data.get('totales', {}))
        
        # Terms
        if 'terminos' in data:
            self._add_terms(doc, data['terminos'])
        
        # Footer
        self._add_footer(doc, data)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer if not output_path else output_path)
        buffer.seek(0)
        
        logger.info("Quotation Word document generated successfully")
        return buffer
    
    def _add_header(self, doc: Document, data: Dict):
        """Add document header"""
        # Title
        title = doc.add_paragraph('COTIZACIÓN', style='CustomTitle')
        
        # Document info
        info = doc.add_paragraph()
        info.add_run(f"No. Cotización: ").bold = True
        info.add_run(f"{data.get('numero', 'N/A')}\n")
        info.add_run(f"Fecha: ").bold = True
        info.add_run(f"{data.get('fecha', datetime.now().strftime('%d/%m/%Y'))}\n")
        info.add_run(f"Válida hasta: ").bold = True
        info.add_run(f"{data.get('valida_hasta', 'N/A')}")
        
        doc.add_paragraph()  # Spacer
    
    def _add_client_info(self, doc: Document, cliente: Dict):
        """Add client information section"""
        # Section heading
        heading = doc.add_paragraph('DATOS DEL CLIENTE', style='CustomHeading')
        
        # Client table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Data
        fields = [
            ('Nombre:', cliente.get('nombre', 'N/A')),
            ('Empresa:', cliente.get('empresa', 'N/A')),
            ('Email:', cliente.get('email', 'N/A')),
            ('Teléfono:', cliente.get('telefono', 'N/A')),
            ('Dirección:', cliente.get('direccion', 'N/A')),
        ]
        
        for idx, (label, value) in enumerate(fields):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
        
        doc.add_paragraph()  # Spacer
    
    def _add_items_table(self, doc: Document, items: List[Dict]):
        """Add items table"""
        # Section heading
        heading = doc.add_paragraph('DETALLE DE SERVICIOS', style='CustomHeading')
        
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ['#', 'Descripción', 'Cantidad', 'Precio Unit.', 'Subtotal']
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            header_cells[idx].paragraphs[0].runs[0].bold = True
        
        # Add items
        for idx, item in enumerate(items, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.get('descripcion', '')
            row.cells[2].text = str(item.get('cantidad', 0))
            row.cells[3].text = f"${item.get('precio_unitario', 0):,.2f}"
            row.cells[4].text = f"${item.get('subtotal', 0):,.2f}"
        
        doc.add_paragraph()  # Spacer
    
    def _add_totals(self, doc: Document, totales: Dict):
        """Add totals section"""
        # Create table for totals
        table = doc.add_table(rows=3, cols=2)
        
        # Data
        totals_data = [
            ('Subtotal:', f"${totales.get('subtotal', 0):,.2f}"),
            ('IVA (16%):', f"${totales.get('iva', 0):,.2f}"),
            ('TOTAL:', f"${totales.get('total', 0):,.2f}"),
        ]
        
        for idx, (label, value) in enumerate(totals_data):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
            
            # Make total row larger and blue
            if idx == 2:
                for run in row.cells[0].paragraphs[0].runs:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(30, 64, 175)
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(30, 64, 175)
        
        doc.add_paragraph()  # Spacer
    
    def _add_terms(self, doc: Document, terminos: str):
        """Add terms and conditions"""
        heading = doc.add_paragraph('TÉRMINOS Y CONDICIONES', style='CustomHeading')
        doc.add_paragraph(terminos)
    
    def _add_footer(self, doc: Document, data: Dict):
        """Add document footer"""
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = footer.add_run(
            f"\nGenerado por PILi Quarts - {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
            f"{data.get('empresa_nombre', 'PILi Engineering')} | {data.get('empresa_contacto', 'contacto@pili.com')}"
        )
        footer_text.font.size = Pt(9)
        footer_text.font.color.rgb = RGBColor(107, 114, 128)
    
    def generate_informe(
        self,
        data: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> BytesIO:
        """
        Generate technical report Word document.
        
        Following python-patterns: Reusable pattern
        """
        logger.info("Generating technical report Word document")
        
        doc = Document()
        self._setup_styles(doc)
        
        # Title
        doc.add_paragraph(data.get('titulo', 'INFORME TÉCNICO'), style='CustomTitle')
        
        # Sections
        for section in data.get('secciones', []):
            doc.add_paragraph(section['titulo'], style='CustomHeading')
            doc.add_paragraph(section['contenido'])
        
        # Save
        buffer = BytesIO()
        doc.save(buffer if not output_path else output_path)
        buffer.seek(0)
        
        logger.info("Technical report Word document generated successfully")
        return buffer
