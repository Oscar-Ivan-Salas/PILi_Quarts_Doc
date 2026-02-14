# -*- coding: utf-8 -*-
"""
Generador de Proyecto Simple - COMPLETO
Proyecto con cronograma, fases y recursos
100% fidelidad al HTML profesional aprobado
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ProyectoSimpleGenerator(BaseDocumentGenerator):
    """Generador para proyectos simples con cronograma"""
    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('PLAN DE PROYECTO')
        run_titulo.font.size = Pt(28)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        nombre_proyecto = self.datos.get('nombre_proyecto', 'Proyecto de Instalación Eléctrica')
        p_nombre = self.doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nombre = p_nombre.add_run(nombre_proyecto)
        run_nombre.font.size = Pt(18)
        run_nombre.font.color.rgb = self.COLOR_SECUNDARIO
        
        codigo = self.datos.get('codigo_proyecto', 'PROY-001')
        p_codigo = self.doc.add_paragraph()
        p_codigo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_codigo = p_codigo.add_run(f'CÓDIGO: {codigo}')
        run_codigo.font.size = Pt(16)
        run_codigo.font.color.rgb = self.COLOR_SECUNDARIO
        run_codigo.font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_info_grid(self):
        """Agrega grid de información clave (4 cards)"""
        # Extraer nombre del cliente si es un diccionario
        cliente_data = self.datos.get('cliente', '')
        if isinstance(cliente_data, dict):
            cliente = cliente_data.get('nombre', '')
        else:
            cliente = str(cliente_data) if cliente_data else ''
        
        # ✅ CORREGIDO: Leer de cronograma anidado primero, fallback a raíz
        cronograma = self.datos.get('cronograma', {})
        duracion = cronograma.get('duracion_total', self.datos.get('duracion_total', '45 días'))
        fecha_inicio = cronograma.get('fecha_inicio', self.datos.get('fecha_inicio', '01/01/2025'))
        fecha_fin = cronograma.get('fecha_fin', self.datos.get('fecha_fin', '28/02/2025'))
        
        # Extraer solo número si viene como "60 días"
        if isinstance(duracion, str):
            duracion_num = ''.join(filter(str.isdigit, duracion)) or '45'
            duracion = duracion_num
        
        # Crear tabla 1x4 para simular grid
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        cards = [
            ('Cliente', cliente),
            ('Duración', f'{duracion} días'),
            ('Inicio', fecha_inicio),
            ('Fin Estimado', fecha_fin)
        ]
        
        for idx, (label, value) in enumerate(cards):
            cell = table.rows[0].cells[idx]
            cell.text = ''
            
            # Label
            p_label = cell.add_paragraph()
            run_label = p_label.add_run(label.upper())
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = RGBColor(107, 114, 128)
            run_label.font.bold = True
            
            # Value
            p_value = cell.add_paragraph()
            run_value = p_value.add_run(str(value))
            run_value.font.size = Pt(14)
            run_value.font.color.rgb = self.COLOR_PRIMARIO
            run_value.font.bold = True
            
            # Fondo claro
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_presupuesto(self):
        """Agrega presupuesto destacado"""
        presupuesto = self.datos.get('presupuesto', '25,000')
        moneda = self.datos.get('moneda', 'PEN')
        
        # Obtener símbolo de moneda
        simbolos_moneda = {
            'PEN': 'S/',
            'USD': '$',
            'EUR': '€',
            'GBP': '£'
        }
        simbolo = simbolos_moneda.get(moneda, 'S/')
        
        p_label = self.doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.add_run('PRESUPUESTO ESTIMADO')
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = RGBColor(107, 114, 128)
        
        p_valor = self.doc.add_paragraph()
        p_valor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_valor = p_valor.add_run(f'{simbolo} {presupuesto:,}' if isinstance(presupuesto, (int, float)) else f'{simbolo} {presupuesto}')
        run_valor.font.size = Pt(32)
        run_valor.font.color.rgb = self.COLOR_PRIMARIO
        run_valor.font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_alcance(self):
        """Agrega alcance del proyecto"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ALCANCE DEL PROYECTO')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        alcance = self.datos.get('alcance_proyecto', 'Descripción del alcance del proyecto...')
        p_alcance = self.doc.add_paragraph(alcance)
        p_alcance.runs[0].font.size = Pt(12)
        p_alcance.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        
        self.doc.add_paragraph()
    
    def _agregar_fases(self):
        """Agrega fases del proyecto (5 fases detalladas)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('FASES DEL PROYECTO')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        fases = self.datos.get('fases', [
            {'numero': 1, 'nombre': 'Inicio y Planificación', 'duracion': 5, 
             'actividades': ['Levantamiento de información', 'Elaboración de propuesta técnica'], 
             'entregable': 'Plan de proyecto aprobado'},
            {'numero': 2, 'nombre': 'Ingeniería y Diseño', 'duracion': 10, 
             'actividades': ['Cálculos técnicos', 'Elaboración de planos'], 
             'entregable': 'Expediente técnico'},
            {'numero': 3, 'nombre': 'Ejecución', 'duracion': 20, 
             'actividades': ['Adquisición de materiales', 'Instalación y montaje'], 
             'entregable': 'Obra ejecutada'},
            {'numero': 4, 'nombre': 'Pruebas y Puesta en Marcha', 'duracion': 5, 
             'actividades': ['Pruebas de funcionamiento', 'Ajustes'], 
             'entregable': 'Sistema operativo'},
            {'numero': 5, 'nombre': 'Cierre', 'duracion': 3, 
             'actividades': ['Documentación as-built', 'Entrega de garantías'], 
             'entregable': 'Proyecto cerrado'}
        ])
        
        for idx, fase in enumerate(fases, 1):
            # Título de fase - usar índice si no hay campo 'numero'
            numero_fase = fase.get('numero', idx)
            p_fase = self.doc.add_paragraph()
            run_num = p_fase.add_run(f"{numero_fase}. {fase['nombre']}")
            run_num.font.size = Pt(16)
            run_num.font.bold = True
            run_num.font.color.rgb = self.COLOR_PRIMARIO
            
            # Duración - extraer número si es string "X días"
            duracion = fase.get('duracion', '0')
            if isinstance(duracion, str):
                # Extraer número de "5 días" -> 5
                duracion_num = ''.join(filter(str.isdigit, duracion)) or '0'
            else:
                duracion_num = str(duracion)
            
            run_dur = p_fase.add_run(f" ({duracion_num} días)")
            run_dur.font.size = Pt(12)
            run_dur.font.color.rgb = RGBColor(255, 255, 255)
            run_dur.font.bold = True
            
            # Actividades - usar descripción si no hay actividades
            actividades = fase.get('actividades', [])
            if not actividades and 'descripcion' in fase:
                # Si no hay actividades pero hay descripción, usarla
                actividades = [fase['descripcion']]
            
            if actividades:
                p_act_titulo = self.doc.add_paragraph()
                run_act_titulo = p_act_titulo.add_run('Actividades:')
                run_act_titulo.font.size = Pt(13)
                run_act_titulo.font.color.rgb = self.COLOR_SECUNDARIO
                run_act_titulo.font.bold = True
                
                for actividad in actividades:
                    p_act = self.doc.add_paragraph(f'▶ {actividad}')
                    p_act.runs[0].font.size = Pt(12)
                    p_act.runs[0].font.color.rgb = RGBColor(55, 65, 81)
                    p_act.paragraph_format.left_indent = Inches(0.25)
            
            # Entregable - usar responsable si no hay entregable
            entregable = fase.get('entregable', fase.get('responsable', ''))
            if entregable:
                p_ent = self.doc.add_paragraph()
                run_ent_label = p_ent.add_run('Entregable: ')
                run_ent_label.font.size = Pt(12)
                run_ent_label.font.bold = True
                run_ent_label.font.color.rgb = self.COLOR_PRIMARIO
                
                run_ent_value = p_ent.add_run(entregable)
                run_ent_value.font.size = Pt(12)
                run_ent_value.font.color.rgb = RGBColor(55, 65, 81)
            
            self.doc.add_paragraph()
    
    def _agregar_recursos(self):
        """Agrega recursos asignados (grid 4 cards)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('RECURSOS ASIGNADOS')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # ✅ CORREGIDO: Aceptar recursos como dict {humanos: [], materiales: []} o array
        recursos_data = self.datos.get('recursos', {})
        if isinstance(recursos_data, dict):
            # Convertir de {humanos: [strings], materiales: [strings]} a formato grid
            humanos = recursos_data.get('humanos', [])
            recursos = []
            for h in humanos[:4]:  # Máximo 4 para grid 2x2
                recursos.append({
                    'rol': h,
                    'cantidad': 1,
                    'dedicacion': '100%',
                    'responsabilidad': '-'
                })
        elif isinstance(recursos_data, list):
            # Ya está en formato correcto
            recursos = recursos_data
        else:
            # Fallback a valores por defecto
            recursos = [
                {'rol': 'Jefe de Proyecto', 'cantidad': 1, 'dedicacion': '25%', 'responsabilidad': 'Coordinación general'},
                {'rol': 'Ingeniero Residente', 'cantidad': 1, 'dedicacion': '100%', 'responsabilidad': 'Ejecución técnica'},
                {'rol': 'Técnicos Instaladores', 'cantidad': 3, 'dedicacion': '100%', 'responsabilidad': 'Instalación'},
                {'rol': 'Inspector de Calidad', 'cantidad': 1, 'dedicacion': '50%', 'responsabilidad': 'Control de calidad'}
            ]
        
        # Crear tabla 2x2 para grid
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        for idx, recurso in enumerate(recursos[:4]):  # Máximo 4
            row = idx // 2
            col = idx % 2
            cell = table.rows[row].cells[col]
            cell.text = ''
            
            # Rol
            p_rol = cell.add_paragraph()
            run_rol = p_rol.add_run(recurso['rol'])
            run_rol.font.size = Pt(14)
            run_rol.font.color.rgb = self.COLOR_PRIMARIO
            run_rol.font.bold = True
            
            # Detalles
            p_det = cell.add_paragraph()
            p_det.add_run(f"Cantidad: {recurso['cantidad']}\n").font.size = Pt(11)
            p_det.add_run(f"Dedicación: {recurso['dedicacion']}\n").font.size = Pt(11)
            p_det.add_run(f"Responsabilidad: {recurso['responsabilidad']}").font.size = Pt(11)
            
            # Fondo claro
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_riesgos(self):
        """Agrega análisis de riesgos (tabla con badges)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ANÁLISIS DE RIESGOS')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        riesgos = self.datos.get('riesgos', [
            {'descripcion': 'Retrasos en entrega de materiales', 'probabilidad': 'Media', 'impacto': 'Alto', 'mitigacion': 'Compra anticipada'},
            {'descripcion': 'Condiciones climáticas adversas', 'probabilidad': 'Baja', 'impacto': 'Medio', 'mitigacion': 'Programación flexible'},
            {'descripcion': 'Cambios en alcance', 'probabilidad': 'Media', 'impacto': 'Alto', 'mitigacion': 'Control de cambios'}
        ])
        
        if riesgos:
            table = self.doc.add_table(rows=len(riesgos) + 1, cols=4)
            table.style = 'Table Grid'
            
            # Header
            headers = ['RIESGO', 'PROBABILIDAD', 'IMPACTO', 'MITIGACIÓN']
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(header)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                shading_elm = OxmlElement('w:shd')
                color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
                shading_elm.set(qn('w:fill'), color_hex)
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Riesgos
            for idx, riesgo in enumerate(riesgos, 1):
                row = table.rows[idx]
                row.cells[0].text = riesgo['descripcion']
                row.cells[1].text = riesgo['probabilidad']
                row.cells[2].text = riesgo['impacto']
                row.cells[3].text = riesgo['mitigacion']
                
                # Centrar probabilidad e impacto
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
    
    def _agregar_entregables(self):
        """Agrega entregables principales (grid con iconos)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ENTREGABLES PRINCIPALES')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        entregables = [
            '📋 Plan de proyecto',
            '📐 Planos de instalación',
            '📄 Especificaciones técnicas',
            '🧮 Memoria de cálculo',
            '✅ Protocolos de pruebas',
            '🏗️ Planos as-built',
            '🛡️ Certificados de garantía'
        ]
        
        # Crear tabla 3x3 para grid
        rows_needed = (len(entregables) + 2) // 3
        table = self.doc.add_table(rows=rows_needed, cols=3)
        table.style = 'Table Grid'
        
        for idx, entregable in enumerate(entregables):
            row = idx // 3
            col = idx % 3
            cell = table.rows[row].cells[col]
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(entregable)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(55, 65, 81)
            run.font.bold = True
            
            # Fondo claro
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_normativa(self):
        """Agrega normativa aplicable"""
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        
        p_label = self.doc.add_paragraph()
        run_label = p_label.add_run('NORMATIVA APLICABLE')
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = RGBColor(107, 114, 128)
        
        p_valor = self.doc.add_paragraph()
        run_valor = p_valor.add_run(normativa)
        run_valor.font.size = Pt(14)
        run_valor.font.color.rgb = self.COLOR_PRIMARIO
        run_valor.font.bold = True
        
        self.doc.add_paragraph()
    
    def generar(self, ruta_salida):
        """Genera el documento completo"""
        self._agregar_header_basico()
        self._agregar_titulo()
        
        # NUEVAS SECCIONES
        self._agregar_info_grid()
        self._agregar_presupuesto()
        self._agregar_alcance()
        
        self._agregar_fases()
        
        # NUEVAS SECCIONES
        self._agregar_recursos()
        self._agregar_riesgos()
        self._agregar_entregables()
        self._agregar_normativa()
        
        self._agregar_footer_basico()
        
        self.doc.save(str(ruta_salida))
        return ruta_salida


def generar_proyecto_simple(datos, ruta_salida, opciones=None):
    """Función de entrada para generar proyecto simple"""
    generator = ProyectoSimpleGenerator(datos, opciones)
    return generator.generar(ruta_salida)
