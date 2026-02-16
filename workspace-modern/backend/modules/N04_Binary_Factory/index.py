import logging
import base64
import json
import os
from pathlib import Path
from pydantic import ValidationError
from .models import BinaryFactoryInput
from .excel_generator import generate_excel_document  # Assuming this exists or will be refactored
# from .word_generator import generate_word_document # Dynamic Import
# from .pdf_generator import generate_pdf_document # Dynamic Import

logger = logging.getLogger("N04_Binary_Factory")

class BinaryFactory:
    def process_request(self, input_data: dict) -> dict:
        """
        Procesa una solicitud de generación de documento validando contra contrato.
        """
        try:
            # 1. Validación Estricta (Contract First con Pydantic)
            try:
                validated_input = BinaryFactoryInput(**input_data)
            except ValidationError as e:
                logger.error(f"Contract Violation: {e.errors()}")
                return {"success": False, "error": f"Contract Violation: {e.errors()}"}

            # 2. Desestructuración Segura
            header = validated_input.header
            branding = validated_input.branding
            payload = validated_input.payload
            output_format = validated_input.output_format

            logger.info(f"🏭 Processing Request: DocType {header.document_type} | Service {header.service_id} | Format {output_format}")

            # 3. Despacho (Factory Pattern) - FORCE V10 ENGINE RECALL
            # Bypass _generate_universal_* methods to enforce Strict Naming and V10 Layouts
            if output_format == "XLSX":
                return self._generate_excel(header, branding, payload)
            elif output_format == "DOCX":
                return self._generate_word(header, branding, payload)
            elif output_format == "PDF":
                return self._generate_pdf(header, branding, payload)
            
            return {"success": False, "error": "Unsupported format"}

        except Exception as e:
            logger.error(f"Critical Error in BinaryFactory: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _generate_universal_pdf(self, template_name, branding, payload, header):
        """Universal Engine: Reads mapping.json and generates PDF"""
        try:
            from fpdf import FPDF
            import io
            
            template_dir = Path(__file__).parent / "templates" / template_name
            mapping_path = template_dir / "mapping.json"
            
            if not mapping_path.exists():
                 return {"success": False, "error": f"Template {template_name} not found"}
            
            with open(mapping_path, "r", encoding="utf-8") as f:
                mapping = json.load(f)
                
            w_layout = mapping.get("word_layout", {})
            title_text = w_layout.get("title", "Documento").replace("{{ service_name }}", f"Servicio {header.service_id}")
            
            pdf = FPDF()
            pdf.add_page()
            
            # Logo
            # If b64 logo provided, we could save temp and use it.
            # Simplified for Seal: Text Header
            pdf.set_font("Arial", 'B', 16)
            pdf.set_text_color(204, 0, 0) # Primary red
            pdf.cell(0, 10, "SOLUCIONES ELECTRICAS", 0, 1, 'C')
            
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, title_text, 0, 1, 'C')
            
            pdf.set_font("Arial", '', 10)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 6, f"Cliente: {payload.client_info.get('nombre', 'N/A')}", 0, 1)
            pdf.cell(0, 6, f"Fecha: {payload.client_info.get('fecha', 'N/A')}", 0, 1)
            pdf.ln(10)
            
            # Mirror Logic: Use Excel Columns to define Order
            # 1. Get Columns Definition
            tables = mapping.get("excel_layout", {}).get("tables", [])
            if not tables:
                 return {"success": False, "error": "No tables defined in mapping"}
            
            columns_map = tables[0].get("columns", {})
            # 2. Sort keys by Value (Column Letter)
            # Logic: B->1, C->2...
            # We assume single letters for now or handle simple comparison
            sorted_keys = sorted(columns_map.keys(), key=lambda k: columns_map[k])
            
            # Table
            headers = w_layout.get("headers", [])
            if not headers:
                # Fallback to Keys if no headers
                headers = [k.capitalize() for k in sorted_keys]
                
            # Dynamic Column Widths based on content count
            # Total width approx 190 (A4)
            # Index (1st) = 15, Others distributed
            col_count = len(headers)
            if col_count == 0: col_count = 1
            
            col_w = []
            available_w = 175
            for i in range(col_count):
                if i == 0: col_w.append(15) # Index/Item usually small
                else: col_w.append(available_w / (col_count - 1))
            
            pdf.set_font("Arial", 'B', 8)
            pdf.set_fill_color(240, 240, 240)
            
            # Draw Headers
            for i, h in enumerate(headers):
                w = col_w[i] if i < len(col_w) else 20
                pdf.cell(w, 8, str(h)[:15], 1, 0, 'C', 1)
            pdf.ln()
            
            pdf.set_font("Arial", '', 8)
            for item in payload.items:
                # Dynamic Row
                for i, key in enumerate(sorted_keys):
                    w = col_w[i] if i < len(col_w) else 20
                    
                    # Special Handlers based on key name?
                    # "index" -> calculated from loop?
                    if key == "index":
                        val = str(payload.items.index(item)+1)
                    else:
                        val = item.get(key, "")
                        
                    # Handle Formula
                    if isinstance(val, str) and val.startswith("="):
                        val = "(Calc)"
                        
                    pdf.cell(w, 8, str(val)[:40], 1)
                pdf.ln()

            pdf.ln(5)
            # Only show Total if it exists in 'totals' object and fits the context
            if payload.totals.get('total'):
                 pdf.cell(0, 8, f"TOTAL: {payload.totals.get('total')}", 0, 1, 'R')
            
            # Output
            try:
                pdf_bytes = pdf.output(dest='S').encode('latin-1')
            except:
                pdf_bytes = pdf.output().encode('latin-1') # Fallback
                
            b64_data = base64.b64encode(pdf_bytes).decode('utf-8')
            
            return {
                "success": True, 
                "filename": f"{template_name}_GEN.pdf",
                "file_b64": b64_data,
                "engine": "Universal PDF v1.0"
            }

        except Exception as e:
            logger.error(f"Universal PDF Gen Error: {e}", exc_info=True)
            return {"success": False, "error": str(e)}
    def _generate_universal_word(self, template_name, branding, payload, header):
        """Universal Engine: Reads mapping.json and generates Word (DOCX)"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            template_dir = Path(__file__).parent / "templates" / template_name
            mapping_path = template_dir / "mapping.json"
            
            if not mapping_path.exists():
                 return {"success": False, "error": f"Template {template_name} not found"}
            
            with open(mapping_path, "r", encoding="utf-8") as f:
                mapping = json.load(f)
                
            w_layout = mapping.get("word_layout", {})
            styles_map = mapping.get("styles", {})
            
            doc = Document()
            
            # Title
            title_text = w_layout.get("title", "Documento").replace("{{ service_name }}", f"Servicio {header.service_id}")
            h1 = doc.add_heading(title_text, 0)
            h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Client Info (Simple implementation)
            doc.add_paragraph(f"Cliente: {payload.client_info.get('nombre', 'N/A')}")
            doc.add_paragraph(f"Fecha: {payload.client_info.get('fecha', 'N/A')}")
            
            # Table
            # Mirror Logic: Use Excel Columns to define Order
            # 1. Get Columns Definition
            tables = mapping.get("excel_layout", {}).get("tables", [])
            if not tables:
                 return {"success": False, "error": "No tables defined in mapping"}
            
            columns_map = tables[0].get("columns", {})
            # 2. Sort keys by Value (Column Letter)
            sorted_keys = sorted(columns_map.keys(), key=lambda k: columns_map[k])
            
            headers = w_layout.get("headers", [])
            if not headers: headers = [k.capitalize() for k in sorted_keys]

            items = payload.items
            
            # Verify Column Count Match?
            # if len(headers) != len(sorted_keys):
            #     logger.warning(f"Header count {len(headers)} != Key count {len(sorted_keys)}")
            
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = str(h)
                
            for item in items:
                row_cells = table.add_row().cells
                for i, key in enumerate(sorted_keys):
                    if i < len(row_cells):
                        if key == "index":
                           val = str(items.index(item) + 1)
                        else:
                           val = str(item.get(key, ""))
                        
                        if val.startswith("="): val = "(Calc)"
                        row_cells[i].text = val
            
            # Totals
            doc.add_paragraph("")
            if payload.totals.get('total'):
                p_total = doc.add_paragraph(f"TOTAL: {payload.totals.get('total', 0)}")
                p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Output
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            b64_data = base64.b64encode(output.read()).decode('utf-8')
            
            return {
                "success": True, 
                "filename": f"{template_name}_GEN.docx",
                "file_b64": b64_data,
                "engine": "Universal Word v1.0"
            }

        except Exception as e:
            logger.error(f"Universal Word Gen Error: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _generate_universal_excel(self, template_name, branding, payload, header):
        """Universal Engine: Reads mapping.json and generates Excel"""
        try:
            template_dir = Path(__file__).parent / "templates" / template_name
            mapping_path = template_dir / "mapping.json"
            
            if not mapping_path.exists():
                 return {"success": False, "error": f"Template {template_name} not found"}
            
            with open(mapping_path, "r", encoding="utf-8") as f:
                mapping = json.load(f)
            
            # --- MICRO ENGINE FOR EXCEL ---
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
            import io
            
            wb = Workbook()
            ws = wb.active
            ws.title = mapping.get("excel_layout", {}).get("sheet_name", "Doc")
            
            layout = mapping.get("excel_layout", {})
            styles = mapping.get("styles", {})
            
            # 1. Static Cells
            for cell_coord, value in layout.get("static_cells", {}).items():
                ws[cell_coord] = value
                ws[cell_coord].font = Font(bold=True, size=12)
            
            # 2. Dynamic Cells
            dyn_map = layout.get("dynamic_cells", {})
            
            # Flatten context for simple mapping resolution
            ctx = {
                "client_name": payload.client_info.get("nombre", "N/A"),
                "total_value": payload.totals.get("total", 0),
                "service_name": f"Service {header.service_id}",
                # Add more as needed by mapping
            }
            
            if "client_name" in dyn_map:
                ws[dyn_map["client_name"]] = ctx["client_name"]
            if "total_value" in dyn_map:
                ws[dyn_map["total_value"]] = ctx["total_value"]
            
            # 3. Tables
            for table_def in layout.get("tables", []):
                data_key = table_def.get("data_key") # "items"
                start_row = table_def.get("start_row", 10)
                cols_map = table_def.get("columns", {})
                
                items = payload.items # List of models (dicts)
                
                current_row = start_row
                for idx, item in enumerate(items, 1):
                    # Item is dict
                    row_data = {
                        "index": idx,
                        "description": item.get("descripcion", ""),
                        "quantity": item.get("cantidad", 0),
                        "total": item.get("total", 0),
                        "unit": item.get("unidad", ""),
                        "price": item.get("precio_unitario", 0)
                    }
                    
                    for field, col_letter in cols_map.items():
                        if field in row_data:
                            cell_ref = f"{col_letter}{current_row}"
                            val = row_data[field]
                            
                            # FORMULA ENGINE (Simple)
                            if isinstance(val, str) and val.startswith("="):
                                # Replace placeholders? No, simpliest is direct formula.
                                # But formulas usually need relative row references like =C{row}*D{row}
                                # "Cuadro de Cargas" specifics require verifying "formulas en el Excel funcionen".
                                # If mapping.json defines a column as a formula template?
                                # E.g. "total": "=E{row}*F{row}"
                                # Let's assume the mapping or the data provides the formula pattern.
                                # Or better: allow mapping.json to define a "formula" for a column.
                                ws[cell_ref] = val.replace("{row}", str(current_row))
                            else:
                                ws[cell_ref] = val
                    
                    current_row += 1

            # Output
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            b64_data = base64.b64encode(output.read()).decode('utf-8')
            
            return {
                "success": True, 
                "filename": f"{template_name}_GEN.xlsx",
                "file_b64": b64_data,
                "engine": "Universal v1.0"
            }

        except Exception as e:
            logger.error(f"Universal Gen Error: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _generate_excel(self, header, branding, payload):
        """Delegates to Excel Generator Strategy"""
        try:
            # Map Document Type ID to Internal Generator Method match
            # 1: Cotizacion Simple
            # 2: Cotizacion Compleja
            # 3: Proyecto Simple
            # 4: Proyecto Complejo
            # 5: Informe Tecnico
            # 6: Informe Ejecutivo
            
            doc_type_map = {
                1: "cotizacion_simple",
                2: "cotizacion_compleja",
                3: "proyecto_simple",
                4: "proyecto_complejo",
                5: "informe_tecnico",
                6: "informe_ejecutivo"
            }
            
            mode = doc_type_map.get(header.document_type, "cotizacion_simple")
            
            # Call the Generator function (refactored to accept structured data)
            # The excel_generator needs to be adaptable. 
            # For now, we assume `generate_excel_document` handles the switch internally or we pass the mode.
            
            result = generate_excel_document(
                mode=mode,
                items=payload.items,
                totals=payload.totals,
                branding={
                    "logo": branding.logo_b64,
                    "color": branding.color_hex
                },
                metadata={
                    "service_id": header.service_id,
                    "user_id": header.user_id,
                    "technical_notes": payload.technical_notes,
                    "client_info": payload.client_info
                }
            )
            
            return result
            
    def _generate_word(self, header, branding, payload):
        """Delegates to HTML-to-Word Generator Strategy (The 'Black Box' - HTML Fidelity)"""
        try:
            from .html_to_word_generator import html_to_word_generator
            
            # Map Document Type ID/String to HTML Generator Method
            doc_type_map = {
                1: "cotizacion_simple",
                2: "cotizacion_compleja",
                3: "proyecto_simple",
                4: "proyecto_complejo",
                5: "informe_tecnico",
                6: "informe_ejecutivo",
                "ELECTRICIDAD_COTIZACION_SIMPLE": "cotizacion_simple",
                "ELECTRICIDAD_COT_COMPLEJA": "cotizacion_compleja",
                "ELECTRICIDAD_PROYECTO_SIMPLE": "proyecto_simple",
                "ELECTRICIDAD_PROYECTO_COMPLEJO": "proyecto_complejo",
                "ELECTRICIDAD_INFORME_TECNICO": "informe_tecnico",
                "ELECTRICIDAD_INFORME_EJECUTIVO": "informe_ejecutivo"
            }
            
            mode = doc_type_map.get(header.document_type)
            if not mode:
                # String check fallback
                str_type = str(header.document_type)
                if "COTIZACION_SIMPLE" in str_type: mode = "cotizacion_simple"
                elif "COT_COMPLEJA" in str_type or "COTIZACION_COMPLEJA" in str_type: mode = "cotizacion_compleja"
                elif "PROYECTO_SIMPLE" in str_type: mode = "proyecto_simple"
                elif "PROYECTO_COMPLEJO" in str_type: mode = "proyecto_complejo"
                elif "INFORME_TECNICO" in str_type: mode = "informe_tecnico"
                elif "INFORME_EJECUTIVO" in str_type: mode = "informe_ejecutivo"
                else:
                    return {"success": False, "error": f"No HTML template mapping for {header.document_type}"}

            # Prepare Input Data for HTML Injection
            # Flatten payload for Jinja2/BeautifulSoup
            input_data = {
                "numero": f"DOC-{header.user_id}-{header.service_id}",
                "cliente": payload.client_info,
                "proyecto": f"Proyecto Serv.{header.service_id}",
                "fecha": payload.client_info.get("fecha", ""),
                "servicio_nombre": f"Servicio {header.service_id}",
                "items": payload.items,
                "subtotal": payload.totals.get("subtotal", 0),
                "igv": payload.totals.get("igv", 0),
                "total": payload.totals.get("total", 0),
                "branding_color": branding.color_hex,
                "technical_notes": payload.technical_notes,
                "user_id": header.user_id
            }
            
            # STRICT NAMING CONVENTION: INFORME_TECNICO_0001_TESLA.DOCX
            type_map = {
                "1": "COTIZACION_SIMPLE", "2": "COTIZACION_COMPLEJA",
                "3": "PROYECTO_SIMPLE", "4": "PROYECTO_COMPLEJO",
                "5": "INFORME_TECNICO", "6": "INFORME_EJECUTIVO",
                "ELECTRICIDAD_COTIZACION_SIMPLE": "COTIZACION_SIMPLE",
                "ELECTRICIDAD_COT_COMPLEJA": "COTIZACION_COMPLEJA",
                "ELECTRICIDAD_PROYECTO_SIMPLE": "PROYECTO_SIMPLE",
                "ELECTRICIDAD_PROYECTO_COMPLEJO": "PROYECTO_COMPLEJO",
                "ELECTRICIDAD_INFORME_TECNICO": "INFORME_TECNICO",
                "ELECTRICIDAD_INFORME_EJECUTIVO": "INFORME_EJECUTIVO"
            }
            doc_label = type_map.get(str(header.document_type), f"DOC_{header.document_type}")
            svc_id_fmt = str(header.service_id).zfill(4)
            final_name = f"{doc_label}_{svc_id_fmt}_TESLA.docx"
            
            import tempfile
            from pathlib import Path
            tmp_dir = Path(tempfile.gettempdir())
            output_path = tmp_dir / final_name
            
            # Dispatch to HTML-to-Word (The Logic User Built)
            if mode == "cotizacion_simple":
                path = html_to_word_generator.generar_cotizacion_simple(input_data, ruta_salida=output_path)
            elif mode == "cotizacion_compleja":
                path = html_to_word_generator.generar_cotizacion_compleja(input_data, ruta_salida=output_path)
            elif mode == "proyecto_simple":
                path = html_to_word_generator.generar_proyecto_simple(input_data, ruta_salida=output_path)
            elif mode == "proyecto_complejo":
                path = html_to_word_generator.generar_proyecto_complejo(input_data, ruta_salida=output_path)
            elif mode == "informe_tecnico":
                path = html_to_word_generator.generar_informe_tecnico(input_data, ruta_salida=output_path)
            elif mode == "informe_ejecutivo":
                path = html_to_word_generator.generar_informe_ejecutivo(input_data, ruta_salida=output_path)
            else:
                 path = html_to_word_generator.generar_cotizacion_simple(input_data, ruta_salida=output_path)

            # Read back file
            with open(path, "rb") as f:
                b64_data = base64.b64encode(f.read()).decode('utf-8')
                
            return {
                "success": True, 
                "filename": final_name,
                "file_b64": b64_data,
                "engine": "HTML-to-Word (Restored)"
            }

        except Exception as e:
            logger.error(f"HTML-DOCX Generation Failed: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _generate_pdf(self, header, branding, payload):
        """Delegates to Playwright PDF Strategy (The HTML Mirror)"""
        try:
            # We must use the HTML -> PDF generator (Playwright) as originally intended
            # This logic captures the exact HTML look.
            from .generators.html_to_pdf_generator import generate_pdf_playwright
            
            input_data = {
                "numero": f"DOC-{header.user_id}-{header.service_id}",
                "cliente": payload.client_info.get("nombre", "Cliente"),
                "fecha": payload.client_info.get("fecha", ""),
                "items": payload.items,
                "subtotal": payload.totals.get("subtotal", 0),
                "igv": payload.totals.get("igv", 0),
                "total": payload.totals.get("total", 0),
                "branding": {
                    "logo_b64": branding.logo_b64,
                    "color": branding.color_hex
                },
                "document_type": header.document_type,
                "technical_notes": payload.technical_notes,
                "client_info": payload.client_info # Pass full info
            }
            
            # STRICT NAMING CONVENTION
            type_map = {
                "1": "COTIZACION_SIMPLE", "2": "COTIZACION_COMPLEJA",
                "3": "PROYECTO_SIMPLE", "4": "PROYECTO_COMPLEJO",
                "5": "INFORME_TECNICO", "6": "INFORME_EJECUTIVO",
                "ELECTRICIDAD_COTIZACION_SIMPLE": "COTIZACION_SIMPLE",
                "ELECTRICIDAD_COT_COMPLEJA": "COTIZACION_COMPLEJA",
                "ELECTRICIDAD_PROYECTO_SIMPLE": "PROYECTO_SIMPLE",
                "ELECTRICIDAD_PROYECTO_COMPLEJO": "PROYECTO_COMPLEJO",
                "ELECTRICIDAD_INFORME_TECNICO": "INFORME_TECNICO",
                "ELECTRICIDAD_INFORME_EJECUTIVO": "INFORME_EJECUTIVO"
            }
            doc_label = type_map.get(str(header.document_type), f"DOC_{header.document_type}")
            svc_id_fmt = str(header.service_id).zfill(4)
            final_name = f"{doc_label}_{svc_id_fmt}_TESLA.pdf"
            
            import tempfile
            from pathlib import Path
            tmp_dir = Path(tempfile.gettempdir())
            output_path = tmp_dir / final_name
            
            # Execute Generator
            path = generate_pdf_playwright(input_data, output_path)
            
            # Read and return B64
            with open(path, "rb") as f:
                b64_data = base64.b64encode(f.read()).decode('utf-8')
                
            return {
                "success": True, 
                "filename": final_name,
                "file_b64": b64_data,
                "engine": "Playwright V10 (HTML Mirror)"
            }

        except Exception as e:
            logger.error(f"Playwright PDF Generation Failed: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

# Singleton Instance
binary_factory = BinaryFactory()
