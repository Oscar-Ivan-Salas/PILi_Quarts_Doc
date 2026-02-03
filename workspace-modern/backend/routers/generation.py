"""
Generation Router - API Endpoints for Document Generation
Handles PDF and Word export functionality
"""
from fastapi import APIRouter, HTTPException, Response
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import Optional, Dict, Any
import io
import os
from datetime import datetime

# Import generators
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    print("WeasyPrint not available (GTK3 missing?). Will try ReportLab.")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Fallback PDF generator
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from utils.template_generator import generate_html_template

router = APIRouter(prefix="/api/generate", tags=["generation"])

class GenerateRequest(BaseModel):
    title: str
    type: str
    data: Dict[str, Any]
    color_scheme: Optional[str] = 'azul-tesla'
    font: Optional[str] = 'Calibri'

@router.post("/pdf")
async def generate_pdf(request: GenerateRequest):
    """Generate PDF document from data"""
    try:
        pdf_file = io.BytesIO()
        
        # Strategy 1: WeasyPrint (High Quality)
        if WEASYPRINT_AVAILABLE:
            try:
                html_content = generate_html_template(
                    request.type, 
                    request.data, 
                    request.color_scheme,
                    request.font
                )
                HTML(string=html_content).write_pdf(pdf_file)
                pdf_file.seek(0)
            except Exception as wp_error:
                print(f"WeasyPrint failed: {wp_error}. Falling back...")
                if not REPORTLAB_AVAILABLE:
                    raise wp_error
                # Trigger fallback manually
                raise OSError("WeasyPrint failed")
                
        # Strategy 2: ReportLab (Reliable Fallback)
        elif REPORTLAB_AVAILABLE:
            c = canvas.Canvas(pdf_file, pagesize=A4)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 800, request.title or "Documento Generado")
            
            c.setFont("Helvetica", 12)
            y = 750
            c.drawString(50, y, f"Tipo: {request.type}")
            y -= 20
            c.drawString(50, y, f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
            y -= 40
            
            # Simple content dump
            c.drawString(50, y, "Datos del Documento:")
            y -= 25
            
            # Iterate safely through basic data
            text_object = c.beginText(50, y)
            text_object.setFont("Helvetica", 10)
            
            if request.data.get('cliente'):
                text_object.textLine("CLIENTE:")
                for k, v in request.data['cliente'].items():
                    text_object.textLine(f" - {k}: {v}")
                text_object.textLine(" ")

            if request.data.get('proyecto'):
                text_object.textLine("PROYECTO:")
                for k, v in request.data['proyecto'].items():
                    text_object.textLine(f" - {k}: {v}")
            
            c.drawText(text_object)
            c.showPage()
            c.save()
            pdf_file.seek(0)
            
        else:
             raise HTTPException(
                status_code=500, 
                detail="No PDF generation libraries available (install weasyprint or reportlab)"
            )

        # 3. Return as downloadable file
        filename = f"{request.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
        return Response(
            content=pdf_file.getvalue(), 
            headers=headers, 
            media_type="application/pdf"
        )
        
    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating PDF: {str(e)}")

@router.post("/word")
async def generate_word(request: GenerateRequest):
    """Generate Word document from data"""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500, 
            detail="Word generation library (python-docx) not installed"
        )
        
    try:
        # Create Word document
        doc = Document()
        
        # Add basic content (simplified for now, mimicking structure)
        doc.add_heading(request.title, 0)
        
        p = doc.add_paragraph()
        p.add_run(f"Tipo: {request.type}\n").bold = True
        p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")
        
        # Add data dump for now (robust logic would go in a utility)
        doc.add_heading('Datos del Documento', level=1)
        
        # Client info
        if 'cliente' in request.data:
            doc.add_heading('Cliente', level=2)
            for key, value in request.data['cliente'].items():
                doc.add_paragraph(f"{key.capitalize()}: {value}")
                
        # Project info
        if 'proyecto' in request.data:
            doc.add_heading('Proyecto', level=2)
            for key, value in request.data['proyecto'].items():
                doc.add_paragraph(f"{key.capitalize()}: {value}")
        
        # Save to buffer
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        filename = f"{request.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
        return Response(
            content=docx_file.getvalue(),
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print(f"Error generating Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating Word doc: {str(e)}")
