# -*- coding: utf-8 -*-
"""
Generador de Informe Ejecutivo APA - COMPLETO
Informe ejecutivo profesional con formato APA
100% fidelidad al HTML profesional aprobado
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class InformeEjecutivoAPAGenerator(BaseDocumentGenerator):
    """Generador para informes ejecutivos con formato APA"""
    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('INFORME EJECUTIVO')
        run_titulo.font.size = Pt(30)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        p_formato = self.doc.add_paragraph()
        p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_formato = p_formato.add_run('Formato APA')
        run_formato.font.size = Pt(14)
        run_formato.font.color.rgb = self.COLOR_SECUNDARIO
        run_formato.font.italic = True
        
        titulo = self.datos.get('titulo', 'Informe Ejecutivo')
        p_subtitulo = self.doc.add_paragraph()
        p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_subtitulo = p_subtitulo.add_run(titulo)
        run_subtitulo.font.size = Pt(16)
        run_subtitulo.font.color.rgb = self.COLOR_SECUNDARIO
        run_subtitulo.font.bold = True
        
        autor = self.datos.get('autor', 'Tesla Electricidad')
        p_autor = self.doc.add_paragraph()
        p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_autor = p_autor.add_run(f'Autor: {autor}')
        run_autor.font.size = Pt(13)
        run_autor.font.color.rgb = RGBColor(55, 65, 81)
        
        fecha = self.datos.get('fecha', '01/01/2025')
        codigo = self.datos.get('codigo', 'INF-EJEC-001')
        p_info = self.doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_info = p_info.add_run(f'{fecha} | Código: {codigo}')
        run_info.font.size = Pt(12)
        run_info.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()
    
    def _agregar_info_general(self):
        """Agrega información general"""
        cliente_data = self.datos.get('cliente', {})
        cliente = cliente_data.get('nombre', 'Cliente') if isinstance(cliente_data, dict) else str(cliente_data)
        fecha = self.datos.get('fecha', '01/01/2025')
        codigo = self.datos.get('codigo', 'INF-EJEC-001')
        
        p_cliente = self.doc.add_paragraph()
        run_label = p_cliente.add_run('Cliente: ')
        run_label.font.size = Pt(12)
        run_label.font.bold = True
        run_label.font.color.rgb = self.COLOR_SECUNDARIO
        
        run_value = p_cliente.add_run(cliente)
        run_value.font.size = Pt(12)
        
        p_fecha = self.doc.add_paragraph()
        run_fecha_label = p_fecha.add_run('Fecha: ')
        run_fecha_label.font.size = Pt(12)
        run_fecha_label.font.bold = True
        run_fecha_label.font.color.rgb = self.COLOR_SECUNDARIO
        
        run_fecha_value = p_fecha.add_run(fecha)
        run_fecha_value.font.size = Pt(12)
        
        p_codigo = self.doc.add_paragraph()
        run_codigo_label = p_codigo.add_run('Código: ')
        run_codigo_label.font.size = Pt(12)
        run_codigo_label.font.bold = True
        run_codigo_label.font.color.rgb = self.COLOR_SECUNDARIO
        
        run_codigo_value = p_codigo.add_run(codigo)
        run_codigo_value.font.size = Pt(12)
        
        self.doc.add_paragraph()
    
    def _agregar_abstract(self):
        """Agrega abstract"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ABSTRACT')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        abstract = self.datos.get('abstract', 'Abstract del informe ejecutivo...')
        p_abstract = self.doc.add_paragraph(abstract)
        p_abstract.runs[0].font.size = Pt(12)
        p_abstract.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_abstract.runs[0].font.italic = True
        p_abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p_keywords = self.doc.add_paragraph()
        run_keywords = p_keywords.add_run('Palabras clave: Instalaciones eléctricas, Ingeniería, Automatización')
        run_keywords.font.size = Pt(10)
        run_keywords.font.color.rgb = RGBColor(107, 114, 128)
        run_keywords.font.italic = True
        
        self.doc.add_paragraph()
    
    def _agregar_introduccion(self):
        """Agrega introducción"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('1. INTRODUCCIÓN')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        cliente_data = self.datos.get('cliente', {})
        cliente = cliente_data.get('nombre', 'Cliente') if isinstance(cliente_data, dict) else str(cliente_data)
        
        intro_text = f'El presente informe ejecutivo presenta los hallazgos y recomendaciones derivados del análisis realizado para {cliente}. El estudio se enfoca en proporcionar información técnica relevante para la toma de decisiones estratégicas.'
        
        p_intro = self.doc.add_paragraph(intro_text)
        p_intro.runs[0].font.size = Pt(12)
        p_intro.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_metodologia(self):
        """Agrega metodología"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('2. METODOLOGÍA')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        metodologia = self.datos.get('metodologia', 'Metodología utilizada en el estudio...')
        p_metodologia = self.doc.add_paragraph(metodologia)
        p_metodologia.runs[0].font.size = Pt(12)
        p_metodologia.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_metodologia.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_resultados(self):
        """Agrega resultados"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('3. RESULTADOS')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        resultados = self.datos.get('resultados', 'Resultados obtenidos del análisis...')
        p_resultados = self.doc.add_paragraph(resultados)
        p_resultados.runs[0].font.size = Pt(12)
        p_resultados.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_resultados.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_discusion(self):
        """Agrega discusión"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('4. DISCUSIÓN')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        discusion = self.datos.get('discusion', 'Discusión de los resultados...')
        p_discusion = self.doc.add_paragraph(discusion)
        p_discusion.runs[0].font.size = Pt(12)
        p_discusion.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_discusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_conclusiones(self):
        """Agrega conclusiones"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('CONCLUSIONES')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        conclusion_text = 'Basándose en los resultados obtenidos y la discusión presentada, se concluye que el análisis realizado proporciona información valiosa para la toma de decisiones estratégicas en el ámbito de las instalaciones eléctricas y automatización.'
        
        p_conclusion = self.doc.add_paragraph(conclusion_text)
        p_conclusion.runs[0].font.size = Pt(12)
        p_conclusion.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        p_conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_referencias(self):
        """Agrega referencias bibliográficas (Formato APA)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('REFERENCIAS')
        run_titulo.font.size = Pt(20)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        referencias = self.datos.get('referencias', [
            {'autor': 'Autor, A.', 'año': '2024', 'titulo': 'Título del artículo', 'fuente': 'Revista Científica'}
        ])
        
        for ref in referencias:
            autor = ref.get('autor', 'Autor, A.')
            año = ref.get('año', '2024')
            titulo = ref.get('titulo', 'Título del artículo')
            fuente = ref.get('fuente', 'Revista Científica')
            
            # Formato APA: Autor, A. (Año). Título. Fuente.
            ref_text = f'{autor} ({año}). {titulo}. {fuente}.'
            
            p_ref = self.doc.add_paragraph(ref_text)
            p_ref.runs[0].font.size = Pt(11)
            p_ref.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            p_ref.paragraph_format.left_indent = Inches(0.5)
            p_ref.paragraph_format.first_line_indent = Inches(-0.5)
        
        self.doc.add_paragraph()
    
    def generar(self, ruta_salida):
        """Genera el documento completo"""
        self._configurar_margenes_apa() # ✅ APA MARGINS
        
        self._agregar_header_basico()
        self._agregar_titulo()
        self._agregar_info_general()
        self._agregar_abstract()
        self._agregar_introduccion()
        self._agregar_metodologia()
        self._agregar_resultados()
        self._agregar_discusion()
        self._agregar_conclusiones()
        self._agregar_referencias()
        
        self._agregar_firma_ingenieria() # ✅ SIGNATURES
        self._agregar_footer_basico()
        
        self.doc.save(str(ruta_salida))
        return ruta_salida


def generar_informe_ejecutivo_apa(datos, ruta_salida, opciones=None):
    """Función de entrada para generar informe ejecutivo APA"""
    generator = InformeEjecutivoAPAGenerator(datos, opciones)
    return generator.generar(ruta_salida)
