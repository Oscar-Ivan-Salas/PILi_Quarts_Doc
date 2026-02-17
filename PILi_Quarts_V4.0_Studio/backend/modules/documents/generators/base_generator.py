# -*- coding: utf-8 -*-
"""
Clase Base para Generadores de Documentos
Contiene funcionalidad compartida por todos los generadores
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path


class BaseDocumentGenerator:
    """Clase base con funcionalidad compartida para todos los generadores"""
    
    # Colores Tesla Azul (por defecto)
    COLOR_PRIMARIO = RGBColor(0, 82, 163)      # #0052A3
    COLOR_SECUNDARIO = RGBColor(30, 64, 175)   # #1E40AF
    COLOR_ACENTO = RGBColor(59, 130, 246)      # #3B82F6
    COLOR_CLARO = RGBColor(239, 246, 255)      # #EFF6FF
    
    def __init__(self, datos, opciones=None):
        """
        Inicializa el generador
        
        Args:
            datos: Diccionario con datos del documento
            opciones: Opciones de personalización (colores, fuente, etc.)
        """
        self.datos = datos
        self.opciones = opciones or {}
        self.doc = Document()
        
        # 🔍 DEBUG: Logging de opciones recibidas
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"🎨 BaseDocumentGenerator.__init__() - Opciones recibidas: {self.opciones}")
        
        # Aplicar esquema de colores personalizado
        self._aplicar_colores()
        
        # Configurar márgenes
        self._configurar_margenes()
    
    def _rgb_to_hex(self, rgb_color):
        """Convierte RGBColor a string hexadecimal"""
        if hasattr(rgb_color, '_color'):
            color_int = rgb_color._color
            r = (color_int >> 16) & 0xFF
            g = (color_int >> 8) & 0xFF
            b = color_int & 0xFF
        else:
            r, g, b = self.color_primario_rgb
        return '{:02X}{:02X}{:02X}'.format(r, g, b)
    
    def _aplicar_colores(self):
        """Aplica esquema de colores según opciones"""
        esquema = self.opciones.get('esquema_colores', 'azul-tesla')
        
        # 🔍 DEBUG: Logging de esquema aplicado
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"🎨 BaseDocumentGenerator._aplicar_colores() - Esquema: {esquema}")
        logger.info(f"🎨 BaseDocumentGenerator._aplicar_colores() - self.opciones: {self.opciones}")
        
        esquemas = {
            'azul-tesla': {
                'primario': (0, 82, 163),
                'secundario': (30, 64, 175),
                'acento': (59, 130, 246),
            },
            'rojo-energia': {
                'primario': (139, 0, 0),
                'secundario': (153, 27, 27),
                'acento': (220, 38, 38),
            },
            'verde-ecologico': {
                'primario': (6, 95, 70),
                'secundario': (4, 120, 87),
                'acento': (16, 185, 129),
            },
            'dorado': {
                'primario': (212, 175, 55),
                'secundario': (184, 134, 11),
                'acento': (255, 215, 0),
            },
            'personalizado': {
                'primario': (139, 92, 246),  # Morado #8B5CF6
                'secundario': (124, 58, 237),  # Morado oscuro #7C3AED
                'acento': (167, 139, 250),  # Morado claro #A78BFA
            },
        }
        
        colores = esquemas.get(esquema, esquemas['azul-tesla'])
        
        # Guardar como tuplas RGB para conversión a hex
        self.color_primario_rgb = colores['primario']
        self.color_secundario_rgb = colores['secundario']
        self.color_acento_rgb = colores['acento']
        
        # Crear objetos RGBColor
        self.COLOR_PRIMARIO = RGBColor(*colores['primario'])
        self.COLOR_SECUNDARIO = RGBColor(*colores['secundario'])
        self.COLOR_ACENTO = RGBColor(*colores['acento'])
    
    def _configurar_margenes(self):
        """Configura márgenes del documento"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
    
    def _agregar_header_basico(self):
        """Agrega header básico con logo y datos de empresa"""
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Columna izquierda: Logo
        cell_logo = table.rows[0].cells[0]
        cell_logo.width = Inches(2.5)
        p_logo = cell_logo.paragraphs[0]
        
        # Intentar cargar logo si existe
        logo_path = self.opciones.get('logo_path') if self.opciones else None
        
        if logo_path and Path(logo_path).exists():
            try:
                run_logo = p_logo.add_run()
                run_logo.add_picture(str(logo_path), width=Inches(2.0))
            except Exception as e:
                run_logo = p_logo.add_run('TESLA')
                run_logo.font.size = Pt(24)
                run_logo.font.bold = True
                run_logo.font.color.rgb = RGBColor(255, 255, 255)
        else:
            run_logo = p_logo.add_run('TESLA')
            run_logo.font.size = Pt(24)
            run_logo.font.bold = True
            run_logo.font.color.rgb = RGBColor(255, 255, 255)
            
            # Fondo con color personalizado
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell_logo._element.get_or_add_tcPr().append(shading_elm)
        
        # Solo mostrar subtítulo si NO hay logo
        if not (logo_path and Path(logo_path).exists()):
            p_subtitle = cell_logo.add_paragraph('Electricidad y Automatización')
            p_subtitle.runs[0].font.size = Pt(8)
            p_subtitle.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        
        # Columna derecha: Datos de empresa
        cell_info = table.rows[0].cells[1]
        cell_info.width = Inches(4.0)
        
        p_empresa = cell_info.paragraphs[0]
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_empresa = p_empresa.add_run('TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C.')
        run_empresa.font.size = Pt(12)
        run_empresa.font.bold = True
        run_empresa.font.color.rgb = self.COLOR_PRIMARIO
        
        cell_info.add_paragraph('RUC: 20601138787').runs[0].font.size = Pt(9)
        cell_info.add_paragraph('Jr. Las Ágatas Mz B Lote 09, Urb. San Carlos, SJL').runs[0].font.size = Pt(9)
        cell_info.add_paragraph('Teléfono: 906 315 961').runs[0].font.size = Pt(9)
        cell_info.add_paragraph('Email: ingenieria.teslaelectricidad@gmail.com').runs[0].font.size = Pt(9)
        
        self.doc.add_paragraph()
    
    def _agregar_footer_basico(self):
        """Agrega pie de página básico"""
        section = self.doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        
        # Border top for footer
        # p = footer.add_paragraph()
        # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        text = "TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C. | RUC: 20601138787 | Tel: 906 315 961"
        run = paragraph.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(107, 114, 128)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page Number (Simple simulation)
        # paragraph.add_run(" | Pág. ")
        # In python-docx, page numbers are complex xml fields. Skipping for now to avoid corruption.

    def _agregar_firma_ingenieria(self):
        """Agrega bloque de firmas de ingeniería"""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Firma 1: Elaborado por
        cell_1 = table.rows[0].cells[0]
        cell_1.width = Inches(3.0)
        p1 = cell_1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_line1 = p1.add_run('___________________________\n')
        run_line1.font.bold = True
        run_name1 = p1.add_run('ELABORADO POR\n')
        run_name1.font.size = Pt(8)
        run_name1.font.bold = True
        p1.add_run('Ingeniero Residente\nCIP: XXXXXX').font.size = Pt(8)
        
        # Firma 2: Revisado por (Oscar Salas)
        cell_2 = table.rows[0].cells[1]
        cell_2.width = Inches(3.0)
        p2 = cell_2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_line2 = p2.add_run('___________________________\n')
        run_line2.font.bold = True
        run_name2 = p2.add_run('REVISADO POR\n')
        run_name2.font.size = Pt(8)
        run_name2.font.bold = True
        run_sign = p2.add_run('ING. OSCAR IVAN SALAS\n')
        run_sign.font.size = Pt(9)
        run_sign.font.bold = True
        run_sign.font.color.rgb = self.COLOR_PRIMARIO
        p2.add_run('Jefe de Proyectos\nCIP: 254890').font.size = Pt(8)

        self.doc.add_paragraph()

    def _configurar_margenes_apa(self):
        """Configura márgenes APA (1 pulgada = 2.54 cm)"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    def generar(self, ruta_salida):
        """
        Método abstracto - debe ser implementado por clases hijas
        """
        raise NotImplementedError("Subclases deben implementar generar()")

