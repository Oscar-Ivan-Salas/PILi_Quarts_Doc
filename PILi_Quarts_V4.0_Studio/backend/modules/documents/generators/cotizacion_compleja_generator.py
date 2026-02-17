# -*- coding: utf-8 -*-
"""
Generador de Cotización Compleja
Cotización con capítulos, subtotales y notas técnicas
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class CotizacionComplejaGenerator(BaseDocumentGenerator):
    """Generador para cotizaciones complejas con capítulos"""
    
    # --- VISUAL HELPERS ---
    def _set_cell_background(self, cell, color_hex):
        """Helper to set background color (shading) for a cell"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_border(self, cell, **kwargs):
        """
        Helper to set cell borders.
        Usage: _set_cell_border(cell, top={"sz": 12, "val": "single", "color": "FF0000"})
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        for border_name, data in kwargs.items():
            border = OxmlElement(f'w:{border_name}')
            for k, v in data.items():
                border.set(qn(f'w:{k}'), str(v))
            tcBorders.append(border)

    def _setup_font(self, run, size_pt, bold=False, color_rgb=None):
        run.font.name = 'Calibri'
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = color_rgb

    # --- SECTIONS ---

    def _agregar_titulo(self):
        """Agrega título en CAJA AZUL CLARO (Estilo HTML: .titulo-documento)"""
        # Usamos una tabla de 1 celda para simular el div con background
        self.doc.add_paragraph() # Spacer
        
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        
        # Background #EFF6FF
        self._set_cell_background(cell, "EFF6FF")
        
        # Border Left 6px #0052A3 -> w:sz is in 1/8 pt. 6px approx 36? Let's say thick.
        # w:sz=24 (3pt)
        self._set_cell_border(cell, left={"val": "single", "sz": "48", "color": "0052A3"})
        
        # Contenido
        p_titulo = cell.paragraphs[0]
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_titulo.add_run('COTIZACIÓN PROFESIONAL')
        self._setup_font(run, 24, True, self.COLOR_PRIMARIO)
        
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run('Versión Completa con Ingeniería de Detalle')
        self._setup_font(run_sub, 12, False, self.COLOR_SECUNDARIO)
        run_sub.font.italic = True
        
        numero = self.datos.get('numero', 'COT-000')
        p_num = cell.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_num = p_num.add_run(f'N° {numero}')
        self._setup_font(run_num, 14, True, self.COLOR_SECUNDARIO)
        
        self.doc.add_paragraph() # Spacer

    def _agregar_info_general(self):
        """Agrega info en GRID (Estilo HTML: .info-section .info-box)"""
        # Tabla madre 1 fila, 2 columnas con espacio
        # DOCX no tiene 'gap', usamos una columna vacía al medio
        table = self.doc.add_table(rows=1, cols=3)
        table.autofit = False
        
        # Anchos relativos: 48% - 4% - 48%
        # A4 width approx 6.5 inches printable
        w_box = Inches(3.1)
        w_gap = Inches(0.2)
        
        table.rows[0].cells[0].width = w_box
        table.rows[0].cells[1].width = w_gap
        table.rows[0].cells[2].width = w_box
        
        # --- BOX 1: CLIENTE ---
        cell_1 = table.rows[0].cells[0]
        self._set_cell_background(cell_1, "F9FAFB") # #F9FAFB
        self._set_cell_border(cell_1, 
                              top={"val": "single", "sz": "4", "color": "DBEAFE"},
                              bottom={"val": "single", "sz": "4", "color": "DBEAFE"},
                              left={"val": "single", "sz": "4", "color": "DBEAFE"},
                              right={"val": "single", "sz": "4", "color": "DBEAFE"})
        
        # Header Box 1
        p_h1 = cell_1.paragraphs[0]
        run_h1 = p_h1.add_run('DATOS DEL CLIENTE')
        self._setup_font(run_h1, 11, True, self.COLOR_PRIMARIO)
        
        # Data Box 1
        cliente = self.datos.get('cliente', {})
        nombre = cliente.get('nombre', 'Cliente') if isinstance(cliente, dict) else str(cliente)
        proyecto = self.datos.get('proyecto', 'Proyecto')
        area = self.datos.get('area_m2', '0')
        
        cliente_lines = [
            ("Cliente:", nombre),
            ("Proyecto:", proyecto),
            ("Área:", f"{area} m²")
        ]
        
        for label, val in cliente_lines:
            p = cell_1.add_paragraph()
            r_l = p.add_run(f"{label} ")
            self._setup_font(r_l, 10, True, self.COLOR_SECUNDARIO)
            r_v = p.add_run(val)
            self._setup_font(r_v, 10, False, RGBColor(55, 65, 81))
            
        # --- GAP (Col 2) ---
        table.rows[0].cells[1].text = ""
        
        # --- BOX 2: COTIZACION ---
        cell_2 = table.rows[0].cells[2]
        self._set_cell_background(cell_2, "F9FAFB")
        self._set_cell_border(cell_2, 
                              top={"val": "single", "sz": "4", "color": "DBEAFE"},
                              bottom={"val": "single", "sz": "4", "color": "DBEAFE"},
                              left={"val": "single", "sz": "4", "color": "DBEAFE"},
                              right={"val": "single", "sz": "4", "color": "DBEAFE"})
        
        p_h2 = cell_2.paragraphs[0]
        run_h2 = p_h2.add_run('DATOS DE LA COTIZACIÓN')
        self._setup_font(run_h2, 11, True, self.COLOR_PRIMARIO)
        
        fecha = self.datos.get('fecha', '01/01/2025')
        vigencia = self.datos.get('vigencia', '30 días')
        servicio = self.datos.get('servicio', 'Servicios Eléctricos')
        
        cot_lines = [
            ("Fecha:", fecha),
            ("Vigencia:", vigencia),
            ("Servicio:", servicio)
        ]
        
        for label, val in cot_lines:
            p = cell_2.add_paragraph()
            r_l = p.add_run(f"{label} ")
            self._setup_font(r_l, 10, True, self.COLOR_SECUNDARIO)
            r_v = p.add_run(val)
            self._setup_font(r_v, 10, False, RGBColor(55, 65, 81))
            
        self.doc.add_paragraph()

    def _agregar_alcance(self):
        """Agrega Alcance en Caja con borde Izquierdo (Estilo HTML: .seccion-completa)"""
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        
        self._set_cell_background(cell, "F9FAFB")
        scale_border = "32" # 4px approx
        self._set_cell_border(cell, left={"val": "single", "sz": scale_border, "color": "0052A3"})
        
        # Titulo
        p_tit = cell.paragraphs[0]
        r_tit = p_tit.add_run("ALCANCE DEL PROYECTO")
        self._setup_font(r_tit, 14, True, self.COLOR_PRIMARIO)
        
        # Descripcion
        descripcion = self.datos.get('descripcion_proyecto', 'Descripción detallada...')
        p_desc = cell.add_paragraph(descripcion)
        self._setup_font(p_desc.runs[0], 10, False, RGBColor(55, 65, 81))
        
        # Incluye
        cell.add_paragraph()
        p_inc = cell.add_paragraph()
        r_inc = p_inc.add_run("INCLUYE:")
        self._setup_font(r_inc, 10, True, self.COLOR_PRIMARIO)
        
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        items = [
            f'Ingeniería de detalle con cálculos según {normativa}',
            'Suministro de materiales de primera calidad',
            'Instalación por personal técnico certificado',
            'Pruebas y puesta en marcha',
            'Documentación técnica completa',
            'Garantía de 12 meses'
        ]
        
        for item in items:
            p = cell.add_paragraph(f"✓ {item}")
            p.paragraph_format.left_indent = Inches(0.2)
            self._setup_font(p.runs[0], 9, False, RGBColor(55, 65, 81))
        
        self.doc.add_paragraph()
        
    def _agregar_capitulos(self):
        """Agrega capítulos con items y subtotales (Estilo HTML: Tabla Azul)"""
        capitulos = self.datos.get('capitulos', [])
        
        if not capitulos:
            items = self.datos.get('items', [])
            if items:
                capitulos = [{'nombre': 'Detalle de la Cotización', 'items': items}]
        
        total_general = 0
        
        for cap_idx, capitulo in enumerate(capitulos, 1):
            p_cap = self.doc.add_paragraph()
            r_cap = p_cap.add_run('Detalle de la Cotización' if cap_idx == 1 else f"CAPÍTULO {cap_idx}")
            self._setup_font(r_cap, 14, True, self.COLOR_PRIMARIO)
            
            # Tabla
            items = capitulo.get('items', [])
            if items:
                table = self.doc.add_table(rows=len(items) + 1, cols=6)
                table.style = 'Table Grid'
                
                # HEADERS: [ITEM, DESCRIPCIÓN, CANT., UNIDAD, P. UNIT., TOTAL]
                # Widths: 8%, 42%, 10%, 10%, 15%, 15%
                # Doc width approx 6.5 inches
                widths = [0.52, 2.73, 0.65, 0.65, 0.97, 0.97]
                
                headers = ['ITEM', 'DESCRIPCIÓN', 'CANT.', 'UNIDAD', 'P. UNIT.', 'TOTAL']
                for idx, (header, w) in enumerate(zip(headers, widths)):
                    cell = table.rows[0].cells[idx]
                    cell.width = Inches(w)
                    self._set_cell_background(cell, "0052A3") # Azul Tesla
                    
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(header)
                    self._setup_font(run, 9, True, RGBColor(255, 255, 255))
                
                # ROWS
                subtotal_cap = 0
                for r_idx, item in enumerate(items, 1):
                    row = table.rows[r_idx]
                    
                    vals = [
                        str(r_idx).zfill(2),
                        item.get('descripcion', ''),
                        f"{float(item.get('cantidad', 0)):.2f}",
                        item.get('unidad', 'und'),
                        f"$ {float(item.get('precio_unitario', 0)):,.2f}",
                        f"$ {float(item.get('cantidad', 0) * item.get('precio_unitario', 0)):,.2f}"
                    ]
                    
                    total_item = float(item.get('cantidad', 0)) * float(item.get('precio_unitario', 0))
                    subtotal_cap += total_item
                    
                    for c_idx, val in enumerate(vals):
                        cell = row.cells[c_idx]
                        cell.width = Inches(widths[c_idx])
                        p = cell.paragraphs[0]
                        # Align Right for numbers (indices 2, 4, 5)
                        if c_idx in [2, 4, 5]: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif c_idx == 3: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = p.add_run(val)
                        self._setup_font(run, 9, False, RGBColor(55, 65, 81))
                
                total_general += subtotal_cap
                self.doc.add_paragraph()
        
        return total_general
    
    def _agregar_totales(self, subtotal):
        """Agrega sección de totales (Estilo HTML: Caja Derecha)"""
        # HTML design is a flexbox right aligned.
        # We can simulate this with a 2-column table where the left column is empty spacer, right column is the totals box.
        
        self.doc.add_paragraph()
        
        # Grid layout: Col 1 (Spacer ~60%), Col 2 (Totals Box ~40%)
        table_main = self.doc.add_table(rows=1, cols=2)
        table_main.autofit = False
        table_main.rows[0].cells[0].width = Inches(3.5)
        table_main.rows[0].cells[1].width = Inches(3.0)
        table_main.rows[0].cells[0].text = "" # Spacer
        
        # Cell for totals
        cell_total = table_main.rows[0].cells[1]
        
        # Inner Table for Totals (3 rows, 2 cols)
        t_inner = cell_total.add_table(rows=3, cols=2)
        # We want to border this inner table? Or the cell?
        # Let's border the inner table or use simple rows
        
        igv = subtotal * 0.18
        total = subtotal + igv
        
        data = [
            ("SUBTOTAL:", f"$ {subtotal:,.2f}", False),
            ("IGV (18%):", f"$ {igv:,.2f}", False),
            ("TOTAL:", f"$ {total:,.2f}", True)
        ]
        
        for idx, (label, val, is_total) in enumerate(data):
            row = t_inner.rows[idx]
            c1 = row.cells[0]
            c2 = row.cells[1]
            
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(label)
            
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2 = p2.add_run(val)
            
            if is_total:
                # Total Row Style (Blue Bg, White Text)
                self._set_cell_background(c1, "0052A3")
                self._set_cell_background(c2, "0052A3")
                self._setup_font(r1, 12, True, RGBColor(255, 255, 255))
                self._setup_font(r2, 12, True, RGBColor(255, 255, 255))
            else:
                # Normal Row (Border bottom)
                self._set_cell_border(c1, bottom={"val": "single", "sz": "4", "color": "DBEAFE"})
                self._set_cell_border(c2, bottom={"val": "single", "sz": "4", "color": "DBEAFE"})
                self._setup_font(r1, 10, True, self.COLOR_SECUNDARIO)
                self._setup_font(r2, 10, True, self.COLOR_PRIMARIO)

        # Border around the box? The HTML has `border: 2px solid #0052A3`.
        # DOCX nested table borders are tricky.
        # Simpler: Just rely on the inner rows or try to frame the cell_total?
        # Let's frame `cell_total` of the outer table
        self._set_cell_border(cell_total, 
                              top={"val": "single", "sz": "12", "color": "0052A3"},
                              bottom={"val": "single", "sz": "12", "color": "0052A3"},
                              left={"val": "single", "sz": "12", "color": "0052A3"},
                              right={"val": "single", "sz": "12", "color": "0052A3"})

        self.doc.add_paragraph()
    
    def _agregar_alcance(self):
        """Agrega sección de Alcance del Proyecto"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ALCANCE DEL PROYECTO')
        run_titulo.font.size = Pt(16)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Descripción
        descripcion = self.datos.get('descripcion_proyecto', 'Descripción detallada del proyecto...')
        p_desc = self.doc.add_paragraph(descripcion)
        p_desc.runs[0].font.size = Pt(11)
        p_desc.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        
        # INCLUYE:
        p_incluye = self.doc.add_paragraph()
        run_incluye = p_incluye.add_run('INCLUYE:')
        run_incluye.font.size = Pt(11)
        run_incluye.font.bold = True
        run_incluye.font.color.rgb = self.COLOR_PRIMARIO
        
        # Lista de items incluidos
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        items_incluidos = [
            f'Ingeniería de detalle con cálculos según {normativa}',
            'Suministro de materiales de primera calidad',
            'Instalación por personal técnico certificado',
            'Pruebas y puesta en marcha',
            'Documentación técnica completa',
            'Garantía de 12 meses'
        ]
        
        for item in items_incluidos:
            p_item = self.doc.add_paragraph(f'✓ {item}')
            p_item.runs[0].font.size = Pt(10)
            p_item.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            p_item.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_paragraph()
    
    def _agregar_cronograma(self):
        """Agrega sección de Cronograma Estimado (Estilo HTML: Grid Azul)"""
        # Spacer
        self.doc.add_paragraph()
        
        # Container Box
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, "F9FAFB")
        self._set_cell_border(cell, left={"val": "single", "sz": "32", "color": "0052A3"})
        
        # Title
        p_tit = cell.paragraphs[0]
        r_tit = p_tit.add_run("CRONOGRAMA ESTIMADO")
        self._setup_font(r_tit, 14, True, self.COLOR_PRIMARIO)
        
        # Inner Grid (4 cols)
        # Nested table inside the cell? Or just a new table below? 
        # Better to add a new table below for the actual cronograma content to avoid nesting complexity
        self.doc.add_paragraph()
        
        cronograma = self.datos.get('cronograma', {})
        fases = [
            {'num': 1, 'nombre': 'Ingeniería', 'dias': cronograma.get('dias_ingenieria', 5)},
            {'num': 2, 'nombre': 'Adquisiciones', 'dias': cronograma.get('dias_adquisiciones', 7)},
            {'num': 3, 'nombre': 'Instalación', 'dias': cronograma.get('dias_instalacion', 10)},
            {'num': 4, 'nombre': 'Pruebas', 'dias': cronograma.get('dias_pruebas', 3)}
        ]
        
        table_crono = self.doc.add_table(rows=2, cols=4)
        table_crono.style = 'Table Grid'
        
        for idx, fase in enumerate(fases):
            # Row 1: Header (Blue)
            cell_h = table_crono.rows[0].cells[idx]
            self._set_cell_background(cell_h, "0052A3")
            p_h = cell_h.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_h = p_h.add_run(str(fase['num']))
            self._setup_font(run_h, 16, True, RGBColor(255, 255, 255))
            
            # Row 2: Content (White)
            cell_c = table_crono.rows[1].cells[idx]
            p_nom = cell_c.paragraphs[0]
            p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_nom = p_nom.add_run(fase['nombre'])
            self._setup_font(run_nom, 10, True, self.COLOR_PRIMARIO)
            
            p_dias = cell_c.add_paragraph()
            p_dias.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_dias = p_dias.add_run(f"{fase['dias']} días")
            self._setup_font(run_dias, 9, False, RGBColor(107, 114, 128))

        self.doc.add_paragraph()

    def _agregar_garantias(self):
        """Agrega sección de Garantías Incluidas (Estilo HTML: Cards)"""
        # Container
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, "F9FAFB")
        self._set_cell_border(cell, left={"val": "single", "sz": "32", "color": "0052A3"})
        
        p_tit = cell.paragraphs[0]
        r_tit = p_tit.add_run("GARANTÍAS INCLUIDAS")
        self._setup_font(r_tit, 14, True, self.COLOR_PRIMARIO)
        
        self.doc.add_paragraph()
        
        garantias_data = self.datos.get('garantias', [])
        garantias = []
        if garantias_data:
            for g in garantias_data:
                garantias.append(g.get('texto', str(g)) if isinstance(g, dict) else str(g))
        else:
            garantias = ['12 meses en mano de obra', 'Garantía de fabricante en equipos', 'Soporte técnico por 6 meses']
            
        # Grid 3 cols
        cols = min(len(garantias), 3)
        if cols == 0: cols = 1
        table_g = self.doc.add_table(rows=1, cols=cols)
        table_g.style = 'Table Grid'
        
        for idx, texto in enumerate(garantias):
            if idx >= cols: break
            cell = table_g.rows[0].cells[idx]
            self._set_cell_background(cell, "FFFFFF") # White card
            
            icons = ["🛠️", "⚙️", "💬"]
            icon = icons[idx % 3]
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_i = p.add_run(icon)
            r_i.font.size = Pt(20)
            
            p_t = cell.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_t = p_t.add_run(texto)
            self._setup_font(r_t, 9, True, RGBColor(55, 65, 81))
            
        self.doc.add_paragraph()

    def _agregar_condiciones_pago(self):
        """Agrega sección de Condiciones de Pago (Estilo HTML: Lista Styled)"""
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, "F9FAFB")
        self._set_cell_border(cell, left={"val": "single", "sz": "32", "color": "0052A3"})
        
        p = cell.paragraphs[0]
        r = p.add_run("CONDICIONES DE PAGO")
        self._setup_font(r, 14, True, self.COLOR_PRIMARIO)
        
        condiciones = self.datos.get('condiciones_pago', [])
        if not condiciones:
             condiciones = ['50% de adelanto a la firma del contrato', '30% al 50% de avance de obra', '20% contra entrega y conformidad']
             
        for cond in condiciones:
            c_text = cond.get('texto', str(cond)) if isinstance(cond, dict) else str(cond)
            p_item = cell.add_paragraph(f"✓ {c_text}")
            p_item.paragraph_format.left_indent = Inches(0.2)
            self._setup_font(p_item.runs[0], 9, False, RGBColor(55, 65, 81))
            
        self.doc.add_paragraph()

    def _agregar_observaciones(self):
        """Agrega observaciones técnicas detalladas (Estilo HTML: Lista Styled)"""
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, "F9FAFB")
        self._set_cell_border(cell, left={"val": "single", "sz": "32", "color": "0052A3"})
        
        p = cell.paragraphs[0]
        r = p.add_run("OBSERVACIONES TÉCNICAS")
        self._setup_font(r, 14, True, self.COLOR_PRIMARIO)
        
        # Content logic same as original but styled
        observaciones_data = self.datos.get('observaciones', [])
        observaciones = []
        if observaciones_data:
             observaciones = [o.get('texto', str(o)) if isinstance(o, dict) else str(o) for o in observaciones_data]
        else:
             normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
             vigencia = self.datos.get('vigencia', '30 días')
             observaciones = [
                f'Trabajos ejecutados según {normativa}',
                'Materiales de primera calidad',
                'Mano de obra especializada',
                'Incluye transporte de materiales',
                'Pruebas y puesta en marcha incluidas',
                'Capacitación al personal del cliente',
                'Documentación técnica completa',
                'Precios en dólares americanos (USD)',
                f'Cotización válida por {vigencia}'
            ]
            
        for obs in observaciones:
            p_item = cell.add_paragraph(f"✓ {obs}")
            p_item.paragraph_format.left_indent = Inches(0.2)
            self._setup_font(p_item.runs[0], 9, False, RGBColor(55, 65, 81))
            
    def _agregar_footer_basico(self):
        """Agrega pie de página con borde superior (Estilo HTML: .footer)"""
        self.doc.add_paragraph()
        
        # Table 1 cell with top border
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._set_cell_border(cell, top={"val": "single", "sz": "24", "color": "0052A3"})
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C.')
        self._setup_font(r, 10, True, self.COLOR_PRIMARIO)
        
        lines = [
            'RUC: 20601138787 | Teléfono: 906 315 961',
            'Email: ingenieria.teslaelectricidad@gmail.com | Jr. Las Ágatas Mz B Lote 09, SJL'
        ]
        
        for line in lines:
            p_l = cell.add_paragraph(line)
            p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._setup_font(p_l.runs[0], 8, False, RGBColor(107, 114, 128))

    def generar(self, ruta_salida):
        """Genera el documento completo con proteccion de fallos"""
        import logging
        logger = logging.getLogger(__name__)
        logger.info("🚀 INICIANDO GENERACIÓN COTIZACIÓN (MODO NATIVO DE ALTA FIDELIDAD)...")
        
        # 1. PRE-CÁLCULO DE SUBTOTAL (Para no romper el orden visual)
        subtotal = 0
        try:
            items = self.datos.get('items', [])
            capitulos = self.datos.get('capitulos', [])
            
            # Sumar items sueltos
            for item in items:
                cant = float(item.get('cantidad', 0))
                precio = float(item.get('precio_unitario', 0))
                subtotal += cant * precio
            
            # Sumar items de capítulos
            for cap in capitulos:
                for item in cap.get('items', []):
                    cant = float(item.get('cantidad', 0))
                    precio = float(item.get('precio_unitario', 0))
                    subtotal += cant * precio
                    
            logger.info(f"💰 Subtotal Pre-calculado: {subtotal}")
        except Exception as e:
            logger.error(f"❌ Error calculando subtotal: {e}")
            subtotal = 0

        # 2. DEFINICIÓN DE ORDEN VISUAL ESTRICTO
        # Cada método añade contenido al documento EN ORDEN.
        methods = [
            (self._agregar_header_basico, "Header Básico"),
            (self._agregar_titulo, "Título"),
            (self._agregar_info_general, "Info General"),
            (self._agregar_alcance, "Alcance"),
            (self._agregar_capitulos, "Capítulos"),  # Ahora sí se llama en orden
            (lambda: self._agregar_totales(subtotal), "Totales"), # Pasamos el subtotal pre-calculado
            (self._agregar_cronograma, "Cronograma"),
            (self._agregar_garantias, "Garantías"),
            (self._agregar_condiciones_pago, "Condiciones Pago"),
            (self._agregar_observaciones, "Observaciones"),
            (self._agregar_footer_basico, "Footer")
        ]

        # 3. EJECUCIÓN SECUENCIAL
        for method, name in methods:
            try:
                logger.info(f"👉 Ejecutando: {name}")
                method()
            except Exception as e:
                logger.error(f"❌ ERROR EN {name.upper()}: {e}", exc_info=True)
                p_error = self.doc.add_paragraph()
                run_error = p_error.add_run(f"[ERROR GENERANDO SECCIÓN {name}: {e}]")
                run_error.font.color.rgb = RGBColor(255, 0, 0)
                run_error.font.bold = True
        
        try:
            logger.info(f"💾 Guardando documento en: {ruta_salida}")
            self.doc.save(str(ruta_salida))
            return ruta_salida
        except Exception as e:
            logger.error(f"❌ ERROR FINAL GUARDANDO ARCHIVO: {e}", exc_info=True)
            raise e



def generar_cotizacion_compleja(datos, ruta_salida, opciones=None):
    """Función de entrada para generar cotización compleja"""
    generator = CotizacionComplejaGenerator(datos, opciones)
    return generator.generar(ruta_salida)
