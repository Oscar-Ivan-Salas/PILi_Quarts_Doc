# -*- coding: utf-8 -*-
"""
Generador Profesional de Cotizaciones Simples
Basado en PLANTILLA_HTML_COTIZACION_SIMPLE.html

Genera documentos Word con diseño profesional que coincide EXACTAMENTE
con la vista previa HTML.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path


class CotizacionSimpleGenerator:
    """Generador de cotizaciones simples con diseño profesional"""
    
    # Colores Tesla Azul (por defecto)
    COLOR_PRIMARIO = RGBColor(0, 82, 163)      # #0052A3
    COLOR_SECUNDARIO = RGBColor(30, 64, 175)   # #1E40AF
    COLOR_ACENTO = RGBColor(59, 130, 246)      # #3B82F6
    COLOR_CLARO = RGBColor(239, 246, 255)      # #EFF6FF
    
    def __init__(self, datos, opciones=None):
        """
        Inicializa el generador
        
        Args:
            datos: Diccionario con datos de la cotización
            opciones: Opciones de personalización (colores, fuente, etc.)
        """
        self.datos = datos
        self.opciones = opciones or {}
        self.doc = Document()
        
        # Aplicar esquema de colores personalizado
        self._aplicar_colores()
        
        # Configurar márgenes
        self._configurar_margenes()
    
    def _rgb_to_hex(self, rgb_color):
        """Convierte RGBColor a string hexadecimal"""
        # Extraer valores RGB del objeto RGBColor
        # RGBColor almacena los valores internamente, necesitamos acceder correctamente
        if hasattr(rgb_color, '_color'):
            # Formato interno de python-docx
            color_int = rgb_color._color
            r = (color_int >> 16) & 0xFF
            g = (color_int >> 8) & 0xFF
            b = color_int & 0xFF
        else:
            # Usar valores almacenados en tupla
            r, g, b = self.color_primario_rgb
        return '{:02X}{:02X}{:02X}'.format(r, g, b)
    
    def _aplicar_colores(self):
        """Aplica esquema de colores según opciones"""
        esquema = self.opciones.get('esquema_colores', 'azul-tesla')
        
        esquemas = {
            'azul-tesla': {
                'primario': (0, 82, 163),
                'secundario': (30, 64, 175),
                'acento': (59, 130, 246),
            },
            'rojo-energia': {
                'primario': (139, 0, 0),
                'secundario': (153, 27, 27),
                'acento': (220, 38, 38),
            },
            'verde-ecologico': {
                'primario': (6, 95, 70),
                'secundario': (4, 120, 87),
                'acento': (16, 185, 129),
            },
            'dorado': {
                'primario': (212, 175, 55),
                'secundario': (184, 134, 11),
                'acento': (255, 215, 0),
            },
            'personalizado': {
                'primario': (147, 51, 234),  # Morado/Lila #9333EA
                'secundario': (126, 34, 206),  # Morado oscuro #7E22CE
                'acento': (168, 85, 247),  # Morado claro #A855F7
            },
        }
        
        colores = esquemas.get(esquema, esquemas['azul-tesla'])
        
        # Guardar como tuplas RGB para conversión a hex
        self.color_primario_rgb = colores['primario']
        self.color_secundario_rgb = colores['secundario']
        self.color_acento_rgb = colores['acento']
        
        # Crear objetos RGBColor
        self.COLOR_PRIMARIO = RGBColor(*colores['primario'])
        self.COLOR_SECUNDARIO = RGBColor(*colores['secundario'])
        self.COLOR_ACENTO = RGBColor(*colores['acento'])
    
    def _configurar_margenes(self):
        """Configura márgenes del documento"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.79)    # 20mm
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(0.79)
            section.right_margin = Inches(0.79)
    
    def _agregar_header(self):
        """Agrega header profesional en la sección de encabezado de Word"""
        # Verificar si se debe mostrar el logo
        if not self.opciones.get('mostrar_logo', True):
            return

        section = self.doc.sections[0]
        header = section.header
        
        # Preparar encabezado (limpiar basura)
        if len(header.paragraphs) > 0:
            header.paragraphs[0].text = ""
            header.paragraphs[0].paragraph_format.space_after = Pt(0)

        # Tabla para layout del header (2 columnas)
        table = header.add_table(rows=1, cols=2, width=Inches(7.2))
        table.autofit = False
        table.allow_autofit = False
        
        # Columna izquierda: Logo
        cell_logo = table.rows[0].cells[0]
        cell_logo.width = Inches(3.2)
        p_logo = cell_logo.paragraphs[0]
        p_logo.paragraph_format.space_after = Pt(0)
        
        # Intentar cargar logo si existe
        logo_path = self.opciones.get('logo_path') if self.opciones else None
        
        if logo_path and Path(logo_path).exists():
            try:
                run_logo = p_logo.add_run()
                # Logo más grande y nítido
                run_logo.add_picture(str(logo_path), width=Inches(2.5))
            except Exception as e:
                run_logo = p_logo.add_run('TESLA')
                run_logo.font.size = Pt(26)
                run_logo.font.bold = True
                run_logo.font.color.rgb = self.COLOR_PRIMARIO
        else:
            # Fallback notable
            run_logo = p_logo.add_run('TESLA ELECTRICIDAD')
            run_logo.font.size = Pt(22)
            run_logo.font.bold = True
            run_logo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Columna derecha: Datos de empresa
        cell_info = table.rows[0].cells[1]
        cell_info.width = Inches(4.0)
        
        p_empresa = cell_info.paragraphs[0]
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_empresa.paragraph_format.space_after = Pt(0)
        run_empresa = p_empresa.add_run('TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C.')
        run_empresa.font.size = Pt(11)
        run_empresa.font.bold = True
        run_empresa.font.color.rgb = self.COLOR_PRIMARIO
        
        # Detalles de empresa simplificados
        info_lines = [
            'RUC: 20601138787',
            'ingenieria.teslaelectricidad@gmail.com'
        ]
        
        for linea in info_lines:
            p = cell_info.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)


    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        # Título principal
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('COTIZACIÓN DE SERVICIOS')
        run_titulo.font.size = Pt(28)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Número de cotización
        numero = self.datos.get('numero', f"COT-{datetime.now().strftime('%Y%m%d%H%M')}")
        p_numero = self.doc.add_paragraph()
        p_numero.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_numero = p_numero.add_run(f'N° {numero}')
        run_numero.font.size = Pt(16)
        run_numero.font.bold = True
        run_numero.font.color.rgb = self.COLOR_SECUNDARIO
        
        self.doc.add_paragraph()
    
    def _agregar_info_general(self):
        """Agrega sección de información general"""
        # Tabla 2x2 para información
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Celda 1: Datos del Cliente
        cell1 = table.rows[0].cells[0]
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run('DATOS DEL CLIENTE')
        run1.font.size = Pt(11)
        run1.font.bold = True
        run1.font.color.rgb = self.COLOR_PRIMARIO
        
        # Extraer datos del cliente (puede ser dict o string)
        cliente_data = self.datos.get('cliente', 'Cliente')
        if isinstance(cliente_data, dict):
            cliente = cliente_data.get('nombre', 'Cliente')
        else:
            cliente = str(cliente_data)
        
        proyecto = self.datos.get('proyecto', 'Proyecto')
        area = self.datos.get('area_m2', '0')
        
        cell1.add_paragraph(f'Cliente: {cliente}').runs[0].font.size = Pt(10)
        cell1.add_paragraph(f'Proyecto: {proyecto}').runs[0].font.size = Pt(10)
        cell1.add_paragraph(f'Área: {area} m²').runs[0].font.size = Pt(10)
        
        # Celda 2: Datos de la Cotización
        cell2 = table.rows[0].cells[1]
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run('DATOS DE LA COTIZACIÓN')
        run2.font.size = Pt(11)
        run2.font.bold = True
        run2.font.color.rgb = self.COLOR_PRIMARIO
        
        fecha = self.datos.get('fecha', datetime.now().strftime('%d/%m/%Y'))
        vigencia = self.datos.get('vigencia', '30 días')
        servicio = self.datos.get('servicio', 'Instalaciones Eléctricas')
        
        cell2.add_paragraph(f'Fecha: {fecha}').runs[0].font.size = Pt(10)
        cell2.add_paragraph(f'Vigencia: {vigencia}').runs[0].font.size = Pt(10)
        cell2.add_paragraph(f'Servicio: {servicio}').runs[0].font.size = Pt(10)
        
        self.doc.add_paragraph()
    
    def _agregar_tabla_items(self):
        """Agrega tabla de items con diseño profesional"""
        # Título de sección
        p_titulo = self.doc.add_paragraph('Detalle de la Cotización')
        p_titulo.runs[0].font.size = Pt(14)
        p_titulo.runs[0].font.bold = True
        p_titulo.runs[0].font.color.rgb = self.COLOR_PRIMARIO
        
        # Tabla de items
        items = self.datos.get('items', [])
        if not items:
            self.doc.add_paragraph('No hay items en esta cotización')
            return
        
        # Crear tabla
        table = self.doc.add_table(rows=1 + len(items), cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['ITEM', 'DESCRIPCIÓN', 'CANT.', 'UNIDAD', 'P. UNIT.', 'TOTAL']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Estilo del header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fondo con color personalizado
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Datos
        for idx, item in enumerate(items, 1):
            row = table.rows[idx]
            
            # Item number
            row.cells[0].text = str(idx).zfill(2)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Descripción
            row.cells[1].text = item.get('descripcion', '')
            
            # Cantidad
            cantidad = item.get('cantidad', 0)
            row.cells[2].text = f"{cantidad:.2f}"
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Unidad
            row.cells[3].text = item.get('unidad', 'und')
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Precio unitario
            precio = item.get('precio_unitario', 0) or item.get('precioUnitario', 0)
            row.cells[4].text = f"S/ {precio:.2f}"
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Total
            total_item = cantidad * precio
            row.cells[5].text = f"S/ {total_item:.2f}"
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[5].paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_totales(self):
        """Agrega sección de totales"""
        # Calcular totales
        items = self.datos.get('items', [])
        subtotal = sum((item.get('cantidad', 0) * (item.get('precio_unitario', 0) or item.get('precioUnitario', 0))) for item in items)
        igv = subtotal * 0.18
        total = subtotal + igv
        
        # Tabla de totales (alineada a la derecha)
        table = self.doc.add_table(rows=3, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Subtotal
        table.rows[0].cells[0].text = 'SUBTOTAL:'
        table.rows[0].cells[1].text = f'S/ {subtotal:.2f}'
        
        # IGV
        table.rows[1].cells[0].text = 'IGV (18%):'
        table.rows[1].cells[1].text = f'S/ {igv:.2f}'
        
        # Total
        table.rows[2].cells[0].text = 'TOTAL:'
        table.rows[2].cells[1].text = f'S/ {total:.2f}'
        
        # Estilo de la última fila (total)
        for cell in table.rows[2].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Fondo con color personalizado
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_observaciones(self):
        """Agrega sección de observaciones"""
        p_titulo = self.doc.add_paragraph('Observaciones Técnicas')
        p_titulo.runs[0].font.size = Pt(12)
        p_titulo.runs[0].font.bold = True
        p_titulo.runs[0].font.color.rgb = self.COLOR_PRIMARIO
        
        observaciones = [
            'Trabajos ejecutados según CNE - Código Nacional de Electricidad',
            'Materiales de primera calidad con certificación',
            'Mano de obra especializada',
            'Garantía de 12 meses en mano de obra',
            'Precios en soles peruanos (PEN)',
            'Forma de pago: 50% adelanto, 50% contra entrega',
            f"Cotización válida por {self.datos.get('vigencia', '30 días')}"
        ]
        
        for obs in observaciones:
            p = self.doc.add_paragraph(obs, style='List Bullet')
            p.runs[0].font.size = Pt(10)
    
    def _agregar_footer(self):
        """Agrega pie de página"""
        self.doc.add_paragraph()
        
        p_footer = self.doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_footer.add_run('TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C.')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_PRIMARIO
        
        contacto = [
            'RUC: 20601138787',
            'Email: ingenieria.teslaelectricidad@gmail.com',
            'Dpto de diseño GatoMichuy huacacayo peru'
        ]
        
        for linea in contacto:
            p = self.doc.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(107, 114, 128)
    
    def generar(self, ruta_salida):
        """
        Genera el documento Word completo
        
        Args:
            ruta_salida: Ruta donde guardar el documento
        
        Returns:
            Ruta del documento generado
        """
        # Construir documento
        self._agregar_header()
        self._agregar_titulo()
        self._agregar_info_general()
        self._agregar_tabla_items()
        self._agregar_totales()
        self._agregar_observaciones()
        self._agregar_footer()
        
        # Guardar
        self.doc.save(ruta_salida)
        return ruta_salida


def generar_cotizacion_simple(datos, ruta_salida, opciones=None):
    """
    Función helper para generar cotización simple
    
    Args:
        datos: Diccionario con datos de la cotización
        ruta_salida: Ruta donde guardar el documento
        opciones: Opciones de personalización
    
    Returns:
        Ruta del documento generado
    """
    generator = CotizacionSimpleGenerator(datos, opciones)
    return generator.generar(ruta_salida)
