"""
Generador de Word V2 - Directo desde JSON sin HTML parsing
Usa python-docx para crear documentos Word profesionales
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Dict, Any
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

class WordGeneratorV2:
    """Generador de documentos Word usando python-docx (sin HTML)"""
    
    # Esquemas de colores profesionales (RGB)
    COLOR_SCHEMES = {
        'azul-tesla': {
            'primary': RGBColor(0, 51, 102),  # Azul oscuro
            'secondary': RGBColor(41, 128, 185),  # Azul claro
            'accent': RGBColor(52, 152, 219)  # Azul brillante
        },
        'rojo-energia': {
            'primary': RGBColor(192, 57, 43),  # Rojo oscuro
            'secondary': RGBColor(231, 76, 60),  # Rojo claro
            'accent': RGBColor(236, 112, 99)  # Rojo brillante
        },
        'verde-ecologico': {
            'primary': RGBColor(39, 174, 96),  # Verde oscuro
            'secondary': RGBColor(46, 204, 113),  # Verde claro
            'accent': RGBColor(88, 214, 141)  # Verde brillante
        },
        'personalizado': {
            'primary': RGBColor(142, 68, 173),  # Púrpura
            'secondary': RGBColor(155, 89, 182),  # Púrpura claro
            'accent': RGBColor(187, 143, 206)  # Púrpura brillante
        }
    }
    
    def __init__(self):
        self.output_dir = Path("storage/generados")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info("✅ WordGeneratorV2 inicializado")
    
    def _get_colors(self, datos: Dict[str, Any]) -> Dict:
        """Obtener colores según esquema seleccionado"""
        personalizacion = datos.get('personalizacion', {})
        esquema = personalizacion.get('esquema_colores', 'azul-tesla')
        return self.COLOR_SCHEMES.get(esquema, self.COLOR_SCHEMES['azul-tesla'])
    
    def _get_font_name(self, datos: Dict[str, Any]) -> str:
        """Obtener nombre de fuente"""
        personalizacion = datos.get('personalizacion', {})
        return personalizacion.get('fuente', 'Calibri')
    
    def _get_font_size(self, datos: Dict[str, Any]) -> int:
        """Obtener tamaño de fuente"""
        personalizacion = datos.get('personalizacion', {})
        return personalizacion.get('tamano_fuente', 11)
    
    def generar_cotizacion(self, datos: Dict[str, Any]) -> Path:
        """
        Generar documento Word desde datos JSON puros
        Funciona para los 6 tipos: cotizacion-simple, cotizacion-compleja,
        proyecto-simple, proyecto-complejo, informe-tecnico, informe-ejecutivo
        
        Args:
            datos: {
                'tipo_documento': str (opcional),
                'cliente': {'nombre': str, 'ruc': str, ...} o str,
                'proyecto': str,
                'numero': str,
                'fecha': str,
                'items': [{'descripcion': str, 'cantidad': float, ...}],
                'subtotal': float,
                'igv': float,
                'total': float
            }
            
        Returns:
            Path al archivo .docx generado
        """
        tipo_doc = datos.get('tipo_documento', 'cotizacion')
        logger.info(f"📄 Generando {tipo_doc} Word V2: {datos.get('numero')}")
        
        
        # 🆕 USAR GENERADOR PROFESIONAL para TODOS los tipos
        tipos_profesionales = [
            'cotizacion-simple', 'cotizacion', 'cotizacion-compleja',
            'proyecto-simple', 'proyecto', 'proyecto-complejo', 'proyecto-pmi',
            'informe-tecnico', 'informe', 'informe-ejecutivo', 'informe-apa'
        ]
        
        if tipo_doc in tipos_profesionales:
            try:
                from .generators import generar_documento
                
                # Preparar opciones de personalización
                personalizacion = datos.get('personalizacion', {})
                
                # Extraer logo si existe (puede venir como base64 o URL)
                logo_path = None
                if personalizacion.get('logo_base64'):
                    # Convertir base64 a archivo temporal
                    import base64
                    import tempfile
                    import os
                    
                    try:
                        logo_base64 = personalizacion.get('logo_base64')
                        
                        # Remover prefijo data:image si existe
                        if ',' in logo_base64:
                            logo_base64 = logo_base64.split(',')[1]
                        
                        # Decodificar base64
                        logo_data = base64.b64decode(logo_base64)
                        
                        # Crear archivo temporal
                        temp_dir = tempfile.gettempdir()
                        temp_filename = f"logo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                        logo_path = os.path.join(temp_dir, temp_filename)
                        
                        # Guardar archivo
                        with open(logo_path, 'wb') as f:
                            f.write(logo_data)
                        
                        logger.info(f"✅ Logo convertido de base64 a: {logo_path}")
                    except Exception as e:
                        logger.error(f"❌ Error convirtiendo logo base64: {e}")
                        logo_path = None
                        
                elif personalizacion.get('logo_url'):
                    logo_path = personalizacion.get('logo_url')
                
                opciones = {
                    'esquema_colores': personalizacion.get('esquema_colores', 'azul-tesla'),
                    'fuente': personalizacion.get('fuente', 'Calibri'),
                    'tamano_fuente': personalizacion.get('tamano_fuente', 11),
                    'mostrar_logo': personalizacion.get('mostrar_logo', True),
                    'logo_path': logo_path
                }
                
                # Generar nombre de archivo
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                numero = datos.get('numero', f'DOC-{timestamp}')
                filename = f"{numero.replace('/', '-')}_{timestamp}.docx"
                output_path = self.output_dir / filename
                
                # Mapear tipo de documento al generador correcto
                tipo_generador = tipo_doc
                if tipo_doc == 'cotizacion':
                    tipo_generador = 'cotizacion-simple'
                elif tipo_doc == 'proyecto':
                    tipo_generador = 'proyecto-simple'
                elif tipo_doc == 'proyecto-pmi':
                    tipo_generador = 'proyecto-complejo'
                elif tipo_doc == 'informe':
                    tipo_generador = 'informe-tecnico'
                elif tipo_doc == 'informe-apa':
                    tipo_generador = 'informe-ejecutivo'
                
                # Usar generador profesional
                logger.info(f"✨ Usando generador profesional para {tipo_doc} → {tipo_generador}")
                generar_documento(tipo_generador, datos, output_path, opciones)
                
                logger.info(f"✅ Documento profesional generado: {output_path}")
                return output_path
                
            except ImportError as e:
                logger.warning(f"⚠️ Generador profesional no disponible: {e}")
                logger.info("🔄 Usando generador V2 clásico como fallback")
                import traceback
                traceback.print_exc()
            except Exception as e:
                logger.error(f"❌ Error en generador profesional: {e}")
                logger.info("🔄 Usando generador V2 clásico como fallback")
                import traceback
                traceback.print_exc()
        
        # Fallback: Usar generador V2 clásico
        
        # Crear documento
        doc = Document()
        
        # Configurar márgenes y estilo
        self._configurar_documento(doc)
        
        # Encabezado (con título dinámico según tipo)
        self._agregar_encabezado(doc, datos)
        
        # Datos del cliente
        self._agregar_datos_cliente(doc, datos)
        
        # Tabla de items
        self._agregar_tabla_items(doc, datos)
        
        # Totales
        self._agregar_totales(doc, datos)
        
        # Pie de página
        self._agregar_pie_pagina(doc, datos)
        
        # Guardar con nombre según tipo
        tipo_prefix = tipo_doc.split('-')[0]  # cotizacion, proyecto, informe
        filename = f"{tipo_prefix}_{datos.get('numero', 'DOC')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        logger.info(f"✅ Word generado: {output_path}")
        return output_path
    
    def _configurar_documento(self, doc):
        """Configurar estilos del documento"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def _agregar_encabezado(self, doc, datos):
        """Agregar encabezado con título dinámico, colores y logo personalizados"""
        tipo_doc = datos.get('tipo_documento', 'cotizacion')
        colors = self._get_colors(datos)
        font_name = self._get_font_name(datos)
        personalizacion = datos.get('personalizacion', {})
        
        # Determinar título según tipo
        if 'proyecto' in tipo_doc:
            titulo_texto = "PROYECTO DE SERVICIOS"
        elif 'informe' in tipo_doc:
            if 'tecnico' in tipo_doc:
                titulo_texto = "INFORME TÉCNICO"
            else:
                titulo_texto = "INFORME EJECUTIVO"
        else:
            titulo_texto = "COTIZACIÓN DE SERVICIOS"
        
        # Logo (si está habilitado y disponible)
        if personalizacion.get('mostrar_logo') and personalizacion.get('logo_base64'):
            try:
                import base64
                import io
                from docx.shared import Inches
                
                # Decodificar base64
                logo_data = personalizacion['logo_base64']
                if ',' in logo_data:
                    logo_data = logo_data.split(',')[1]  # Remover prefijo data:image/...
                
                logo_bytes = base64.b64decode(logo_data)
                logo_stream = io.BytesIO(logo_bytes)
                
                # Determinar alineación según posición
                posicion = personalizacion.get('posicion_logo', 'center')
                if posicion == 'left':
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif posicion == 'right':
                    alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:  # center
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Agregar logo con alineación personalizada
                logo_para = doc.add_paragraph()
                logo_para.alignment = alignment
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_stream, width=Inches(1.5))
                
                logger.info(f"✅ Logo agregado al documento (posición: {posicion})")
            except Exception as e:
                logger.warning(f"⚠️ No se pudo agregar logo: {e}")
        
        # Título con color personalizado
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run(titulo_texto)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = font_name
        run.font.color.rgb = colors['primary']  # Color personalizado
        
        # Número y fecha
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f"N° {datos.get('numero', 'N/A')} | Fecha: {datos.get('fecha', 'N/A')}")
        run.font.name = font_name
        run.font.color.rgb = colors['secondary']
        
        doc.add_paragraph()  # Espacio
    
    def _agregar_datos_cliente(self, doc, datos):
        """Agregar información del cliente con colores personalizados"""
        cliente = datos.get('cliente', {})
        colors = self._get_colors(datos)
        font_name = self._get_font_name(datos)
        font_size = self._get_font_size(datos)
        
        # Extraer nombre de cliente (puede ser dict o string)
        if isinstance(cliente, dict):
            nombre_cliente = cliente.get('nombre', '[Cliente]')
            ruc_cliente = cliente.get('ruc', '')
            direccion_cliente = cliente.get('direccion', '')
            telefono_cliente = cliente.get('telefono', '')
            email_cliente = cliente.get('email', '')
        else:
            nombre_cliente = str(cliente) if cliente else '[Cliente]'
            ruc_cliente = ''
            direccion_cliente = ''
            telefono_cliente = ''
            email_cliente = ''
        
        # Título sección con color personalizado
        p = doc.add_paragraph()
        run = p.add_run("DATOS DEL CLIENTE")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = font_name
        run.font.color.rgb = colors['primary']
        
        # Datos con fuente personalizada
        for text in [
            f"Cliente: {nombre_cliente}",
            f"RUC: {ruc_cliente}" if ruc_cliente else None,
            f"Dirección: {direccion_cliente}" if direccion_cliente else None,
            f"Teléfono: {telefono_cliente}" if telefono_cliente else None,
            f"Email: {email_cliente}" if email_cliente else None
        ]:
            if text:
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
        
        # Proyecto
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("PROYECTO")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = font_name
        run.font.color.rgb = colors['primary']
        
        p = doc.add_paragraph(f"Nombre: {datos.get('proyecto', '[Proyecto]')}")
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        if datos.get('descripcion'):
            p = doc.add_paragraph(f"Descripción: {datos.get('descripcion', '')}")
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
        
        doc.add_paragraph()  # Espacio
    
    def _agregar_tabla_items(self, doc, datos):
        """Agregar tabla de items con colores y opciones de visualización"""
        colors = self._get_colors(datos)
        font_name = self._get_font_name(datos)
        font_size = self._get_font_size(datos)
        personalizacion = datos.get('personalizacion', {})
        ocultar_precios = personalizacion.get('ocultar_precios_unitarios', False)
        
        # Título con color personalizado
        p = doc.add_paragraph()
        run = p.add_run("DETALLE DE LA COTIZACIÓN")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = font_name
        run.font.color.rgb = colors['primary']
        
        # Crear tabla (columnas dinámicas según opciones)
        items = datos.get('items', [])
        
        # 🐛 DEBUG: Ver qué items recibimos
        logger.info(f"🔍 DEBUG TABLA - Total items recibidos: {len(items)}")
        for idx, item in enumerate(items):
            logger.info(f"  Item {idx}: desc='{item.get('descripcion')}', cant={item.get('cantidad')}, precio={item.get('precio_unitario')}")
        
        num_cols = 4 if ocultar_precios else 5
        tabla = doc.add_table(rows=1, cols=num_cols)
        tabla.style = 'Light Grid Accent 1'
        
        # Encabezados con color personalizado
        headers = tabla.rows[0].cells
        col_idx = 0
        headers[col_idx].text = "DESCRIPCIÓN"
        col_idx += 1
        headers[col_idx].text = "CANT."
        col_idx += 1
        headers[col_idx].text = "UNIDAD"
        col_idx += 1
        
        if not ocultar_precios:
            headers[col_idx].text = "P. UNIT."
            col_idx += 1
        
        headers[col_idx].text = "SUBTOTAL"
        
        # Estilo de encabezados
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = font_name
                    run.font.color.rgb = colors['secondary']
        
        # Agregar items
        logger.info(f"📝 Agregando {len(items)} items a la tabla...")
        for idx, item in enumerate(items):
            row = tabla.add_row().cells
            col_idx = 0
            
            descripcion = str(item.get('descripcion', ''))
            cantidad = item.get('cantidad', 0)
            unidad = str(item.get('unidad', 'und'))
            precio_unitario = item.get('precio_unitario', 0)
            
            logger.info(f"  Fila {idx}: {descripcion} | {cantidad} {unidad} | S/ {precio_unitario}")
            
            
            # Agregar datos a las celdas Y aplicar estilo inmediatamente
            # Descripción
            row[col_idx].text = descripcion
            for paragraph in row[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            col_idx += 1
            
            # Cantidad
            row[col_idx].text = str(cantidad)
            for paragraph in row[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            col_idx += 1
            
            # Unidad
            row[col_idx].text = unidad
            for paragraph in row[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            col_idx += 1
            
            # Precio unitario (si no está oculto)
            if not ocultar_precios:
                row[col_idx].text = f"S/ {precio_unitario:,.2f}"
                for paragraph in row[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                col_idx += 1
            
            # Subtotal
            subtotal_item = cantidad * precio_unitario
            row[col_idx].text = f"S/ {subtotal_item:,.2f}"
            for paragraph in row[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            
            logger.info(f"    Subtotal calculado: S/ {subtotal_item:,.2f}")
        
        doc.add_paragraph()  # Espacio
    
    def _agregar_totales(self, doc, datos):
        """Agregar sección de totales con opciones de visualización"""
        colors = self._get_colors(datos)
        font_name = self._get_font_name(datos)
        font_size = self._get_font_size(datos)
        personalizacion = datos.get('personalizacion', {})
        ocultar_igv = personalizacion.get('ocultar_igv', False)
        
        # Crear tabla de totales (filas dinámicas según opciones)
        num_rows = 2 if ocultar_igv else 3
        tabla = doc.add_table(rows=num_rows, cols=2)
        
        row_idx = 0
        
        # Subtotal (siempre mostrar)
        tabla.rows[row_idx].cells[0].text = "Subtotal:"
        tabla.rows[row_idx].cells[1].text = f"S/ {datos.get('subtotal', 0):,.2f}"
        row_idx += 1
        
        # IGV (condicional)
        if not ocultar_igv:
            tabla.rows[row_idx].cells[0].text = "IGV (18%):"
            tabla.rows[row_idx].cells[1].text = f"S/ {datos.get('igv', 0):,.2f}"
            row_idx += 1
        
        # Total
        tabla.rows[row_idx].cells[0].text = "TOTAL:"
        tabla.rows[row_idx].cells[1].text = f"S/ {datos.get('total', 0):,.2f}"
        
        # Aplicar estilos a todas las filas
        for i, row in enumerate(tabla.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        
                        # Última fila (TOTAL) en negrita y color primario
                        if i == num_rows - 1:
                            run.font.bold = True
                            run.font.color.rgb = colors['primary']
                            if j == 1:  # Columna del monto
                                run.font.size = Pt(14)
        
        doc.add_paragraph()  # Espacio
    
    def _agregar_pie_pagina(self, doc, datos):
        """Agregar observaciones y firma"""
        # Observaciones
        p = doc.add_paragraph()
        run = p.add_run("OBSERVACIONES")
        run.font.bold = True
        run.font.size = Pt(10)
        
        doc.add_paragraph(datos.get('observaciones', 'Precios incluyen IGV'))
        doc.add_paragraph(f"Vigencia: {datos.get('vigencia', '30 días')}")

# Instancia global
word_generator_v2 = WordGeneratorV2()
