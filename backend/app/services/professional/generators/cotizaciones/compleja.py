# -*- coding: utf-8 -*-
"""
Generador de Cotización Compleja
Cotización con capítulos, subtotales y notas técnicas
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class CotizacionComplejaGenerator(BaseDocumentGenerator):
    """Generador para cotizaciones complejas con capítulos"""
    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('COTIZACIÓN DE SERVICIOS')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        numero = self.datos.get('numero', 'COT-000')
        p_numero = self.doc.add_paragraph()
        p_numero.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_numero = p_numero.add_run(f'N° {numero}')
        run_numero.font.size = Pt(14)
        run_numero.font.color.rgb = self.COLOR_SECUNDARIO
        
        self.doc.add_paragraph()
    
    def _agregar_info_general(self):
        """Agrega información general en tabla"""
        table = self.doc.add_table(rows=1, cols=2)
        
        # Cliente
        cell_cliente = table.rows[0].cells[0]
        p1 = cell_cliente.paragraphs[0]
        run1 = p1.add_run('DATOS DEL CLIENTE')
        run1.font.size = Pt(11)
        run1.font.bold = True
        run1.font.color.rgb = self.COLOR_PRIMARIO
        
        cliente_data = self.datos.get('cliente', 'Cliente')
        if isinstance(cliente_data, dict):
            cliente = cliente_data.get('nombre', 'Cliente')
        else:
            cliente = str(cliente_data)
        
        proyecto = self.datos.get('proyecto', 'Proyecto')
        area = self.datos.get('area_m2', '0')
        
        cell_cliente.add_paragraph(f'Cliente: {cliente}').runs[0].font.size = Pt(10)
        cell_cliente.add_paragraph(f'Proyecto: {proyecto}').runs[0].font.size = Pt(10)
        cell_cliente.add_paragraph(f'Área: {area} m²').runs[0].font.size = Pt(10)
        
        # Cotización
        cell_cot = table.rows[0].cells[1]
        p2 = cell_cot.paragraphs[0]
        run2 = p2.add_run('DATOS DE LA COTIZACIÓN')
        run2.font.size = Pt(11)
        run2.font.bold = True
        run2.font.color.rgb = self.COLOR_PRIMARIO
        
        fecha = self.datos.get('fecha', '01/01/2025')
        vigencia = self.datos.get('vigencia', '30 días')
        servicio = self.datos.get('servicio', 'Servicios Eléctricos')
        
        cell_cot.add_paragraph(f'Fecha: {fecha}').runs[0].font.size = Pt(10)
        cell_cot.add_paragraph(f'Vigencia: {vigencia}').runs[0].font.size = Pt(10)
        cell_cot.add_paragraph(f'Servicio: {servicio}').runs[0].font.size = Pt(10)
        
        self.doc.add_paragraph()
    
    def _agregar_capitulos(self):
        """Agrega capítulos con items y subtotales"""
        capitulos = self.datos.get('capitulos', [])
        
        if not capitulos:
            # Si no hay capítulos, usar items directamente
            items = self.datos.get('items', [])
            if items:
                capitulos = [{'nombre': 'Detalle de la Cotización', 'items': items}]
        
        total_general = 0
        
        for cap_idx, capitulo in enumerate(capitulos, 1):
            # Título del capítulo
            p_cap = self.doc.add_paragraph()
            run_cap = p_cap.add_run(f"CAPÍTULO {cap_idx}: {capitulo.get('nombre', 'Sin nombre')}")
            run_cap.font.size = Pt(14)
            run_cap.font.bold = True
            run_cap.font.color.rgb = self.COLOR_PRIMARIO
            
            # Tabla de items del capítulo
            items = capitulo.get('items', [])
            if items:
                table = self.doc.add_table(rows=len(items) + 1, cols=6)
                table.style = 'Table Grid'
                
                # Header
                headers = ['ITEM', 'DESCRIPCIÓN', 'CANT.', 'UNIDAD', 'P. UNIT.', 'TOTAL']
                for idx, header in enumerate(headers):
                    cell = table.rows[0].cells[idx]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(header)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Fondo con color personalizado
                    shading_elm = OxmlElement('w:shd')
                    color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
                    shading_elm.set(qn('w:fill'), color_hex)
                    cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Items
                subtotal_cap = 0
                for idx, item in enumerate(items, 1):
                    row = table.rows[idx]
                    cantidad = float(item.get('cantidad', 0))
                    precio = float(item.get('precio_unitario', 0))
                    total_item = cantidad * precio
                    subtotal_cap += total_item
                    
                    row.cells[0].text = str(idx)
                    row.cells[1].text = item.get('descripcion', '')
                    row.cells[2].text = str(cantidad)
                    row.cells[3].text = item.get('unidad', 'und')
                    row.cells[4].text = f"S/ {precio:.2f}"
                    row.cells[5].text = f"S/ {total_item:.2f}"
                
                # Subtotal del capítulo
                p_subtotal = self.doc.add_paragraph()
                p_subtotal.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_subtotal = p_subtotal.add_run(f"Subtotal Capítulo {cap_idx}: S/ {subtotal_cap:.2f}")
                run_subtotal.font.size = Pt(12)
                run_subtotal.font.bold = True
                run_subtotal.font.color.rgb = self.COLOR_SECUNDARIO
                
                total_general += subtotal_cap
                self.doc.add_paragraph()
        
        return total_general
    
    def _agregar_totales(self, subtotal):
        """Agrega sección de totales"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('RESUMEN DE COSTOS')
        run_titulo.font.size = Pt(14)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Tabla de totales
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        igv = subtotal * 0.18
        total = subtotal + igv
        
        # Subtotal
        table.rows[0].cells[0].text = 'SUBTOTAL:'
        table.rows[0].cells[1].text = f'S/ {subtotal:.2f}'
        
        # IGV
        table.rows[1].cells[0].text = 'IGV (18%):'
        table.rows[1].cells[1].text = f'S/ {igv:.2f}'
        
        # Total
        cell_total_label = table.rows[2].cells[0]
        cell_total_valor = table.rows[2].cells[1]
        
        p_label = cell_total_label.paragraphs[0]
        run_label = p_label.add_run('TOTAL:')
        run_label.font.bold = True
        run_label.font.size = Pt(14)
        run_label.font.color.rgb = RGBColor(255, 255, 255)
        
        p_valor = cell_total_valor.paragraphs[0]
        run_valor = p_valor.add_run(f'S/ {total:.2f}')
        run_valor.font.bold = True
        run_valor.font.size = Pt(14)
        run_valor.font.color.rgb = RGBColor(255, 255, 255)
        
        # Fondo con color personalizado
        for cell in [cell_total_label, cell_total_valor]:
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_alcance(self):
        """Agrega sección de Alcance del Proyecto"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ALCANCE DEL PROYECTO')
        run_titulo.font.size = Pt(16)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Descripción
        descripcion = self.datos.get('descripcion_proyecto', 'Descripción detallada del proyecto...')
        p_desc = self.doc.add_paragraph(descripcion)
        p_desc.runs[0].font.size = Pt(11)
        p_desc.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        
        # INCLUYE:
        p_incluye = self.doc.add_paragraph()
        run_incluye = p_incluye.add_run('INCLUYE:')
        run_incluye.font.size = Pt(11)
        run_incluye.font.bold = True
        run_incluye.font.color.rgb = self.COLOR_PRIMARIO
        
        # Lista de items incluidos
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        items_incluidos = [
            f'Ingeniería de detalle con cálculos según {normativa}',
            'Suministro de materiales de primera calidad',
            'Instalación por personal técnico certificado',
            'Pruebas y puesta en marcha',
            'Documentación técnica completa',
            'Garantía de 12 meses'
        ]
        
        for item in items_incluidos:
            p_item = self.doc.add_paragraph(f'✓ {item}')
            p_item.runs[0].font.size = Pt(10)
            p_item.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            p_item.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_paragraph()
    
    def _agregar_cronograma(self):
        """Agrega sección de Cronograma Estimado"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('CRONOGRAMA ESTIMADO')
        run_titulo.font.size = Pt(16)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Obtener datos del cronograma
        cronograma = self.datos.get('cronograma', {})
        fases = [
            {'num': 1, 'nombre': 'Ingeniería', 'dias': cronograma.get('dias_ingenieria', 5)},
            {'num': 2, 'nombre': 'Adquisiciones', 'dias': cronograma.get('dias_adquisiciones', 7)},
            {'num': 3, 'nombre': 'Instalación', 'dias': cronograma.get('dias_instalacion', 10)},
            {'num': 4, 'nombre': 'Pruebas', 'dias': cronograma.get('dias_pruebas', 3)}
        ]
        
        # Crear tabla para las fases
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        for idx, fase in enumerate(fases):
            # Fila 1: Número de fase
            cell_num = table.rows[0].cells[idx]
            p_num = cell_num.paragraphs[0]
            p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_num = p_num.add_run(str(fase['num']))
            run_num.font.size = Pt(18)
            run_num.font.bold = True
            run_num.font.color.rgb = RGBColor(255, 255, 255)
            
            # Fondo con color
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell_num._element.get_or_add_tcPr().append(shading_elm)
            
            # Fila 2: Nombre y días
            cell_info = table.rows[1].cells[idx]
            cell_info.text = ''
            
            p_nombre = cell_info.add_paragraph()
            p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_nombre = p_nombre.add_run(fase['nombre'])
            run_nombre.font.size = Pt(11)
            run_nombre.font.bold = True
            run_nombre.font.color.rgb = self.COLOR_PRIMARIO
            
            p_dias = cell_info.add_paragraph()
            p_dias.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_dias = p_dias.add_run(f"{fase['dias']} días")
            run_dias.font.size = Pt(10)
            run_dias.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()
    
    def _agregar_garantias(self):
        """Agrega sección de Garantías Incluidas"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('GARANTÍAS INCLUIDAS')
        run_titulo.font.size = Pt(16)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        garantias = [
            '🛠️ 12 meses en mano de obra',
            '⚙️ Garantía de fabricante en equipos',
            '💬 Soporte técnico por 6 meses'
        ]
        
        # Crear tabla para las garantías
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        for idx, garantia in enumerate(garantias):
            cell = table.rows[0].cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(garantia)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(55, 65, 81)
            run.font.bold = True
            
            # Fondo claro
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_condiciones_pago(self):
        """Agrega sección de Condiciones de Pago"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('CONDICIONES DE PAGO')
        run_titulo.font.size = Pt(16)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        condiciones = [
            '50% de adelanto a la firma del contrato',
            '30% al 50% de avance de obra',
            '20% contra entrega y conformidad'
        ]
        
        for condicion in condiciones:
            p_cond = self.doc.add_paragraph(f'✓ {condicion}')
            p_cond.runs[0].font.size = Pt(10)
            p_cond.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            p_cond.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_paragraph()
    
    def _agregar_observaciones(self):
        """Agrega observaciones técnicas detalladas"""
        p_obs_titulo = self.doc.add_paragraph()
        run_obs_titulo = p_obs_titulo.add_run('OBSERVACIONES TÉCNICAS')
        run_obs_titulo.font.size = Pt(16)
        run_obs_titulo.font.bold = True
        run_obs_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        vigencia = self.datos.get('vigencia', '30 días')
        
        observaciones = [
            f'Trabajos ejecutados según {normativa}',
            'Materiales de primera calidad con certificación internacional',
            'Mano de obra especializada y certificada',
            'Incluye transporte de materiales',
            'Pruebas y puesta en marcha incluidas',
            'Capacitación al personal del cliente',
            'Documentación técnica completa (planos as-built, protocolos, certificados)',
            'Precios en dólares americanos (USD)',
            f'Cotización válida por {vigencia}'
        ]
        
        for obs in observaciones:
            p_obs = self.doc.add_paragraph(f'✓ {obs}')
            p_obs.runs[0].font.size = Pt(10)
            p_obs.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            p_obs.paragraph_format.left_indent = Inches(0.25)
    
    def generar(self, ruta_salida):
        """Genera el documento completo"""
        self._agregar_header_basico()
        self._agregar_titulo()
        self._agregar_info_general()
        
        # Nueva sección: Alcance del Proyecto
        self._agregar_alcance()
        
        subtotal = self._agregar_capitulos()
        self._agregar_totales(subtotal)
        
        # Nuevas secciones: Cronograma, Garantías, Condiciones de Pago
        self._agregar_cronograma()
        self._agregar_garantias()
        self._agregar_condiciones_pago()
        
        # Observaciones mejoradas
        self._agregar_observaciones()
        
        self._agregar_footer_basico()
        
        self.doc.save(str(ruta_salida))
        return ruta_salida


def generar_cotizacion_compleja(datos, ruta_salida, opciones=None):
    """Función de entrada para generar cotización compleja"""
    generator = CotizacionComplejaGenerator(datos, opciones)
    return generator.generar(ruta_salida)
