# -*- coding: utf-8 -*-
"""
Generador de Proyecto Complejo PMI - COMPLETO
Project Charter según metodología PMI con KPIs, Gantt, RACI, Stakeholders
100% fidelidad al HTML profesional aprobado
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ProyectoComplejoPMIGenerator(BaseDocumentGenerator):
    """Generador para proyectos complejos con metodología PMI"""
    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('PROJECT CHARTER')
        run_titulo.font.size = Pt(30)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        p_subtitulo = self.doc.add_paragraph()
        p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_subtitulo = p_subtitulo.add_run('Gestión de Proyectos según Metodología PMI')
        run_subtitulo.font.size = Pt(14)
        run_subtitulo.font.color.rgb = self.COLOR_SECUNDARIO
        run_subtitulo.font.italic = True
        
        nombre_proyecto = self.datos.get('nombre_proyecto', 'SISTEMA DE INSTALACIÓN ELÉCTRICA INDUSTRIAL')
        p_nombre = self.doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nombre = p_nombre.add_run(nombre_proyecto)
        run_nombre.font.size = Pt(18)
        run_nombre.font.color.rgb = self.COLOR_SECUNDARIO
        
        codigo = self.datos.get('codigo_proyecto', 'PROY-PMI-001')
        p_codigo = self.doc.add_paragraph()
        p_codigo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_codigo = p_codigo.add_run(codigo)
        run_codigo.font.size = Pt(14)
        run_codigo.font.color.rgb = self.COLOR_SECUNDARIO
        run_codigo.font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_info_grid(self):
        """Agrega grid de información (4 cards)"""
        # Extraer nombre del cliente si es un diccionario
        cliente_data = self.datos.get('cliente', '')
        if isinstance(cliente_data, dict):
            cliente = cliente_data.get('nombre', '')
        else:
            cliente = str(cliente_data) if cliente_data else ''
        duracion = self.datos.get('duracion_total', 60)
        fecha_inicio = self.datos.get('fecha_inicio', '01/01/2025')
        fecha_fin = self.datos.get('fecha_fin', '28/02/2025')
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        cards = [
            ('Cliente', cliente),
            ('Duración Total', f'{duracion} días'),
            ('Inicio', fecha_inicio),
            ('Fin Estimado', fecha_fin)
        ]
        
        for idx, (label, value) in enumerate(cards):
            cell = table.rows[0].cells[idx]
            cell.text = ''
            
            p_label = cell.add_paragraph()
            run_label = p_label.add_run(label.upper())
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = RGBColor(107, 114, 128)
            run_label.font.bold = True
            
            p_value = cell.add_paragraph()
            run_value = p_value.add_run(str(value))
            run_value.font.size = Pt(14)
            run_value.font.color.rgb = self.COLOR_PRIMARIO
            run_value.font.bold = True
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_presupuesto(self):
        """Agrega presupuesto destacado"""
        presupuesto = self.datos.get('presupuesto', '75,000')
        
        p_label = self.doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.add_run('PRESUPUESTO TOTAL DEL PROYECTO')
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = RGBColor(107, 114, 128)
        
        p_valor = self.doc.add_paragraph()
        p_valor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_valor = p_valor.add_run(f'$ {presupuesto}')
        run_valor.font.size = Pt(36)
        run_valor.font.color.rgb = self.COLOR_PRIMARIO
        run_valor.font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_kpis(self):
        """Agrega KPIs PMI (5 indicadores)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('INDICADORES DE DESEMPEÑO (KPIs)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        spi = self.datos.get('spi', '1.05')
        cpi = self.datos.get('cpi', '0.98')
        ev_k = self.datos.get('ev_k', '45')
        pv_k = self.datos.get('pv_k', '43')
        ac_k = self.datos.get('ac_k', '46')
        
        table = self.doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'
        
        kpis = [
            ('SPI', spi, 'Schedule Performance'),
            ('CPI', cpi, 'Cost Performance'),
            ('EV', f'${ev_k}K', 'Earned Value'),
            ('PV', f'${pv_k}K', 'Planned Value'),
            ('AC', f'${ac_k}K', 'Actual Cost')
        ]
        
        for idx, (label, value, desc) in enumerate(kpis):
            cell_label = table.rows[0].cells[idx]
            cell_label.text = label
            cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_label.paragraphs[0].runs[0].font.size = Pt(12)
            cell_label.paragraphs[0].runs[0].font.bold = True
            
            cell_value = table.rows[1].cells[idx]
            cell_value.text = value
            cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_value.paragraphs[0].runs[0].font.size = Pt(20)
            cell_value.paragraphs[0].runs[0].font.color.rgb = self.COLOR_PRIMARIO
            cell_value.paragraphs[0].runs[0].font.bold = True
            
            cell_desc = table.rows[2].cells[idx]
            cell_desc.text = desc
            cell_desc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_desc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_desc.paragraphs[0].runs[0].font.color.rgb = RGBColor(156, 163, 175)
        
        self.doc.add_paragraph()
    
    def _agregar_alcance(self):
        """Agrega alcance del proyecto (WBS Level 1)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ALCANCE DEL PROYECTO (WBS Level 1)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        alcance = self.datos.get('alcance_proyecto', 'Descripción del alcance del proyecto...')
        p_alcance = self.doc.add_paragraph(alcance)
        p_alcance.runs[0].font.size = Pt(12)
        p_alcance.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_cronograma_gantt(self):
        """Agrega cronograma Gantt"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('CRONOGRAMA DEL PROYECTO (Diagrama Gantt)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        dias_ing = self.datos.get('dias_ingenieria', 15)
        dias_ejec = self.datos.get('dias_ejecucion', 25)
        
        fases = [
            ('1. Inicio y Planificación', '10 días'),
            ('2. Gestión Stakeholders', '3 días'),
            ('3. Ingeniería y Diseño', f'{dias_ing} días'),
            ('4. Ejecución', f'{dias_ejec} días'),
            ('5. Pruebas y Puesta en Marcha', '8 días'),
            ('6. Cierre', '5 días')
        ]
        
        for fase, duracion in fases:
            p_fase = self.doc.add_paragraph()
            run_fase = p_fase.add_run(f'{fase}: {duracion}')
            run_fase.font.size = Pt(11)
            run_fase.font.bold = True
            run_fase.font.color.rgb = RGBColor(55, 65, 81)
        
        self.doc.add_paragraph()
    
    def _agregar_stakeholders(self):
        """Agrega registro de stakeholders"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('REGISTRO DE STAKEHOLDERS')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        stakeholders = self.datos.get('stakeholders', [
            {'nombre': 'Cliente', 'rol': 'Cliente / Patrocinador Principal', 'poder': 'Alto', 'interes': 'Alto'},
            {'nombre': 'Jefe de Proyecto', 'rol': 'Project Manager', 'poder': 'Alto', 'interes': 'Alto'},
            {'nombre': 'Equipo Técnico', 'rol': 'Ingenieros y Técnicos', 'poder': 'Medio', 'interes': 'Alto'}
        ])
        
        for sh in stakeholders:
            p_sh = self.doc.add_paragraph()
            run_nombre = p_sh.add_run(f"{sh['nombre']}: ")
            run_nombre.font.size = Pt(14)
            run_nombre.font.bold = True
            run_nombre.font.color.rgb = self.COLOR_PRIMARIO
            
            run_rol = p_sh.add_run(sh['rol'])
            run_rol.font.size = Pt(11)
            
            p_badges = self.doc.add_paragraph()
            run_badges = p_badges.add_run(f"Poder: {sh['poder']} | Interés: {sh['interes']}")
            run_badges.font.size = Pt(10)
            run_badges.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()
    
    def _agregar_matriz_raci(self):
        """Agrega matriz RACI"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('MATRIZ RACI (Responsabilidades)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        table = self.doc.add_table(rows=6, cols=6)
        table.style = 'Table Grid'
        
        # Header
        headers = ['Actividad', 'PM', 'Ing. Residente', 'Técnicos', 'Inspector QA', 'Cliente']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Actividades
        actividades = [
            ('Planificación del Proyecto', ['A', 'R', 'I', 'C', 'C']),
            ('Diseño e Ingeniería', ['A', 'R', 'C', 'C', 'I']),
            ('Ejecución de Obra', ['A', 'A', 'R', 'C', 'I']),
            ('Control de Calidad', ['A', 'C', 'C', 'R', 'I']),
            ('Aprobación de Entregables', ['R', 'C', 'I', 'C', 'A'])
        ]
        
        for idx, (actividad, roles) in enumerate(actividades, 1):
            table.rows[idx].cells[0].text = actividad
            for j, rol in enumerate(roles, 1):
                cell = table.rows[idx].cells[j]
                cell.text = rol
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
        
        # Leyenda
        p_leyenda = self.doc.add_paragraph()
        run_leyenda = p_leyenda.add_run('Leyenda: R=Responsable | A=Aprobador | C=Consultado | I=Informado')
        run_leyenda.font.size = Pt(10)
        
        self.doc.add_paragraph()
    
    def _agregar_riesgos(self):
        """Agrega registro de riesgos (Top 5)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('REGISTRO DE RIESGOS (Top 5)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        riesgos = self.datos.get('riesgos', [
            {'id': 'R01', 'descripcion': 'Retrasos en entrega de equipos', 'probabilidad': 'Media', 'impacto': 'Alto', 'severidad': 'Alta', 'mitigacion': 'Compra anticipada'},
            {'id': 'R02', 'descripcion': 'Variación de precios', 'probabilidad': 'Media', 'impacto': 'Medio', 'severidad': 'Media', 'mitigacion': 'Precio fijo'},
            {'id': 'R03', 'descripcion': 'Cambios en alcance', 'probabilidad': 'Alta', 'impacto': 'Alto', 'severidad': 'Alta', 'mitigacion': 'Control de cambios'}
        ])
        
        table = self.doc.add_table(rows=len(riesgos) + 1, cols=6)
        table.style = 'Table Grid'
        
        # Header
        headers = ['ID', 'Riesgo', 'Probabilidad', 'Impacto', 'Severidad', 'Plan de Mitigación']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Riesgos
        for idx, riesgo in enumerate(riesgos, 1):
            row = table.rows[idx]
            row.cells[0].text = riesgo['id']
            row.cells[1].text = riesgo['descripcion']
            row.cells[2].text = riesgo['probabilidad']
            row.cells[3].text = riesgo['impacto']
            row.cells[4].text = riesgo['severidad']
            row.cells[5].text = riesgo['mitigacion']
            
            for j in [2, 3, 4]:
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
    
    def _agregar_entregables(self):
        """Agrega entregables principales"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ENTREGABLES PRINCIPALES DEL PROYECTO')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        entregables = [
            '📋 Project Charter', '📊 Plan de gestión', '👥 Registro stakeholders',
            '📝 WBS y diccionario', '📅 Cronograma Gantt', '📐 Planos instalación',
            '📄 Especif. técnicas', '✅ Plan de calidad', '⚠️ Registro de riesgos',
            '🔬 Protocolos FAT/SAT', '🏗️ Planos as-built', '📚 Lecciones aprendidas',
            '🎯 Acta de cierre'
        ]
        
        for entregable in entregables:
            p_ent = self.doc.add_paragraph(f'• {entregable}')
            p_ent.runs[0].font.size = Pt(11)
        
        self.doc.add_paragraph()
    
    def _agregar_normativa(self):
        """Agrega normativa aplicable"""
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        
        p_label = self.doc.add_paragraph()
        run_label = p_label.add_run('NORMATIVA Y ESTÁNDARES APLICABLES')
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = RGBColor(107, 114, 128)
        
        p_valor = self.doc.add_paragraph()
        run_valor = p_valor.add_run(normativa)
        run_valor.font.size = Pt(14)
        run_valor.font.color.rgb = self.COLOR_PRIMARIO
        run_valor.font.bold = True
        
        p_pmbok = self.doc.add_paragraph()
        run_pmbok = p_pmbok.add_run('Gestión según PMBOK® Guide 7th Edition')
        run_pmbok.font.size = Pt(11)
        run_pmbok.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()
    
    def generar(self, ruta_salida):
        """Genera el documento completo"""
        self._agregar_header_basico()
        self._agregar_titulo()
        self._agregar_info_grid()
        self._agregar_presupuesto()
        self._agregar_kpis()
        self._agregar_alcance()
        self._agregar_cronograma_gantt()
        self._agregar_stakeholders()
        self._agregar_matriz_raci()
        self._agregar_riesgos()
        self._agregar_entregables()
        self._agregar_normativa()
        self._agregar_footer_basico()
        
        self.doc.save(str(ruta_salida))
        return ruta_salida


def generar_proyecto_complejo_pmi(datos, ruta_salida, opciones=None):
    """Función de entrada para generar proyecto complejo PMI"""
    generator = ProyectoComplejoPMIGenerator(datos, opciones)
    return generator.generar(ruta_salida)
