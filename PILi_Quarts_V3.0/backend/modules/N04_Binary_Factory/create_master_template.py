
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_master_template():
    doc = Document()
    
    # 1. Header (Logo & Company Name)
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C."
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style.font.bold = True
    paragraph.style.font.size = Pt(14)
    paragraph.style.font.color.rgb = RGBColor(0, 82, 163) # #0052A3
    
    p2 = header.add_paragraph("RUC: 20601138787 | Ingeniera de Detalle y Mantenimiento")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.style.font.size = Pt(9)
    
    # 2. Title (Dynamic)
    doc.add_paragraph("\n")
    h1 = doc.add_heading("{{ TITULO_DOCUMENTO }}", 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. Client Box
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Row 1
    table.cell(0,0).text = "CLIENTE:"
    table.cell(0,1).text = "{{ CLIENTE_NOMBRE }}"
    
    # Row 2
    table.cell(1,0).text = "PROYECTO:"
    table.cell(1,1).text = "{{ PROYECTO_NOMBRE }}"
    
    # Row 3
    table.cell(2,0).text = "FECHA:"
    table.cell(2,1).text = "{{ FECHA }}"
    
    doc.add_paragraph("\n")
    
    # 4. Items Table (Jinja2 Loop for docxtpl)
    doc.add_heading("Detalle de Items", level=2)
    
    table_items = doc.add_table(rows=1, cols=6)
    table_items.style = 'Table Grid'
    hdr_cells = table_items.rows[0].cells
    hdr_cells[0].text = "ITEM"
    hdr_cells[1].text = "DESCRIPCIÓN"
    hdr_cells[2].text = "UNIDAD"
    hdr_cells[3].text = "CANT."
    hdr_cells[4].text = "P.UNIT"
    hdr_cells[5].text = "TOTAL"
    
    # Jinja2 Row
    row_cells = table_items.add_row().cells
    row_cells[0].text = "{{ item.index }}"
    row_cells[1].text = "{{ item.descripcion }}"
    row_cells[2].text = "{{ item.unidad }}"
    row_cells[3].text = "{{ item.cantidad }}"
    row_cells[4].text = "{{ item.precio }}"
    row_cells[5].text = "{{ item.total }}"
    
    # Valid Jinja2 syntax for docxtpl
    # The tr opening/closing tags are handled slightly differently in docxtpl depending on version,
    # but strictly speaking we usually just put the tags in the cells.
    # For a loop, we need {% tr for item in items %} in the first cell? 
    # Or simpler: we insert the tag structure.
    # To be safe for this "Mock Master", we will use the standard jinja loop notation roughly.
    # Actually, let's keep it simple. This file creates the BASE. 
    # docxtpl will need specific tags.
    # Let's add the tags as text.
    
    p = row_cells[0].paragraphs[0]
    p.text = ""
    run = p.add_run("{% for item in items %}")
    run.font.size = Pt(1) # Hide it visually if possible or just put it in the first cell
    p.add_run("{{ item.index }}")
    
    row_cells[5].paragraphs[0].add_run("{% endfor %}")
    
    # 5. Totals
    doc.add_paragraph("\n")
    p_tot = doc.add_paragraph()
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_tot.add_run("SUBTOTAL: {{ SUBTOTAL }}\n").bold = True
    p_tot.add_run("IGV (18%): {{ IGV }}\n").bold = True
    p_tot.add_run("TOTAL: {{ TOTAL }}").bold = True
    p_tot.style.font.size = Pt(11)
    
    # 6. Disclaimer / Footer
    doc.add_section()
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "Documento Generado por Sistema RALFTH - Tesla Electricidad v10.0"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.style.font.size = Pt(8)
    
    # Save
    os.makedirs("e:/PILi_Quarts/workspace-modern/backend/modules/N04_Binary_Factory/templates", exist_ok=True)
    doc.save("e:/PILi_Quarts/workspace-modern/backend/modules/N04_Binary_Factory/templates/master_tesla.docx")
    print("✅ Master Template Created: master_tesla.docx")

if __name__ == "__main__":
    create_master_template()
