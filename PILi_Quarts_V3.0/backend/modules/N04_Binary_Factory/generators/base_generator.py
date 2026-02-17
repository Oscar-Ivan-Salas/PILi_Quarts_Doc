# -*- coding: utf-8 -*-
"""
Clase Base para Generadores de Documentos
Contiene funcionalidad compartida por todos los generadores
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path


class BaseDocumentGenerator:
    """Clase base con funcionalidad compartida para todos los generadores"""
    
    # Colores Tesla Azul (por defecto)
    COLOR_PRIMARIO = RGBColor(0, 82, 163)      # #0052A3
    COLOR_SECUNDARIO = RGBColor(30, 64, 175)   # #1E40AF
    COLOR_ACENTO = RGBColor(59, 130, 246)      # #3B82F6
    COLOR_CLARO = RGBColor(239, 246, 255)      # #EFF6FF
    
    def __init__(self, datos, opciones=None):
        """
        Inicializa el generador
        
        Args:
            datos: Diccionario con datos del documento
            opciones: Opciones de personalizaci칩n (colores, fuente, etc.)
        """
        self.datos = datos
        self.opciones = opciones or {}
        self.doc = Document()
        
        # 游댌 DEBUG: Logging de opciones recibidas
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"游꿛 BaseDocumentGenerator.__init__() - Opciones recibidas: {self.opciones}")
        
        # Aplicar esquema de colores personalizado
        self._aplicar_colores()
        
        # Configurar m치rgenes
        self._configurar_margenes()
    
    def _rgb_to_hex(self, rgb_color):
        """Convierte RGBColor a string hexadecimal"""
        if hasattr(rgb_color, '_color'):
            color_int = rgb_color._color
            r = (color_int >> 16) & 0xFF
            g = (color_int >> 8) & 0xFF
            b = color_int & 0xFF
        else:
            r, g, b = self.color_primario_rgb
        return '{:02X}{:02X}{:02X}'.format(r, g, b)
    
    def _aplicar_colores(self):
        """Aplica esquema de colores seg칰n opciones"""
        esquema = self.opciones.get('esquema_colores', 'azul-tesla')
        
        # 游댌 DEBUG: Logging de esquema aplicado
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"游꿛 BaseDocumentGenerator._aplicar_colores() - Esquema: {esquema}")
        logger.info(f"游꿛 BaseDocumentGenerator._aplicar_colores() - self.opciones: {self.opciones}")
        
        esquemas = {
            'azul-tesla': {
                'primario': (0, 82, 163),
                'secundario': (30, 64, 175),
                'acento': (59, 130, 246),
            },
            'rojo-energia': {
                'primario': (139, 0, 0),
                'secundario': (153, 27, 27),
                'acento': (220, 38, 38),
            },
            'verde-ecologico': {
                'primario': (6, 95, 70),
                'secundario': (4, 120, 87),
                'acento': (16, 185, 129),
            },
            'dorado': {
                'primario': (212, 175, 55),
                'secundario': (184, 134, 11),
                'acento': (255, 215, 0),
            },
            'personalizado': {
                'primario': (139, 92, 246),  # Morado #8B5CF6
                'secundario': (124, 58, 237),  # Morado oscuro #7C3AED
                'acento': (167, 139, 250),  # Morado claro #A78BFA
            },
        }
        
        colores = esquemas.get(esquema, esquemas['azul-tesla'])
        
        # Guardar como tuplas RGB para conversi칩n a hex
        self.color_primario_rgb = colores['primario']
        self.color_secundario_rgb = colores['secundario']
        self.color_acento_rgb = colores['acento']
        
        # Crear objetos RGBColor
        self.COLOR_PRIMARIO = RGBColor(*colores['primario'])
        self.COLOR_SECUNDARIO = RGBColor(*colores['secundario'])
        self.COLOR_ACENTO = RGBColor(*colores['acento'])
    
    def _configurar_margenes(self):
        """Configura m치rgenes del documento"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
    
    def _agregar_header_basico(self):
        """Agrega header profesional en la secci칩n de encabezado de Word"""
        # Verificar si se debe mostrar el logo
        if not self.opciones.get('mostrar_logo', True):
            return

        section = self.doc.sections[0]
        header = section.header
        
        # Limpiar p치rrafos existentes en el header y prepararlos
        if len(header.paragraphs) > 0:
            p_first = header.paragraphs[0]
            # No podemos eliminar el 칰nico p치rrafo de un header f치cil, pero podemos usarlo
            # Sin embargo, add_table suele a침adir un nuevo p치rrafo o bloquearse.
            # La mejor forma es usar el primer p치rrafo para el logo si no usamos tabla, 
            # pero para 2 columnas necesitamos tabla.
            p_first.text = ""
            p_first.paragraph_format.space_after = Pt(0)
            p_first.paragraph_format.line_spacing = 1.0

        # Agregar tabla al header
        table = header.add_table(rows=1, cols=2, width=Inches(7.2))
        table.autofit = False
        table.allow_autofit = False
        
        # Columna izquierda: Logo
        cell_logo = table.rows[0].cells[0]
        cell_logo.width = Inches(3.2)
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Intentar cargar logo si existe
        logo_path = self.opciones.get('logo_path') if self.opciones else None
        
        if logo_path and Path(logo_path).exists():
            try:
                run_logo = p_logo.add_run()
                # Aumentar un poco el tama침o para visibilidad
                run_logo.add_picture(str(logo_path), width=Inches(2.5))
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Error insertando imagen de logo: {e}")
                run_logo = p_logo.add_run('TESLA')
                run_logo.font.size = Pt(26)
                run_logo.font.bold = True
                run_logo.font.color.rgb = self.COLOR_PRIMARIO
        else:
            # Fallback visualmente notable si no hay imagen
            run_logo = p_logo.add_run('TESLA ELECTRICIDAD')
            run_logo.font.size = Pt(22)
            run_logo.font.bold = True
            run_logo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Columna derecha: Datos de empresa
        cell_info = table.rows[0].cells[1]
        cell_info.width = Inches(4.0)
        
        p_empresa = cell_info.paragraphs[0]
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_empresa = p_empresa.add_run('TESLA ELECTRICIDAD Y AUTOMATIZACI칍N S.A.C.')
        run_empresa.font.size = Pt(11)
        run_empresa.font.bold = True
        run_empresa.font.color.rgb = self.COLOR_PRIMARIO
        
        info_lines = [
            'RUC: 20601138787',
            'ingenieria.teslaelectricidad@gmail.com'
        ]
        
        for linea in info_lines:
            p = cell_info.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            p.paragraph_format.space_after = Pt(0)

        # Tratar de eliminar el espacio extra que deja Word bajo la tabla
        # (A veces queda un p치rrafo vac칤o despu칠s de la tabla)


    
    def _agregar_footer_basico(self):
        """Agrega pie de p치gina b치sico"""
        self.doc.add_paragraph()
        
        p_footer = self.doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_footer.add_run('TESLA ELECTRICIDAD Y AUTOMATIZACI칍N S.A.C.')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_PRIMARIO
        
        contacto = [
            'RUC: 20601138787 | Tel칠fono: 906 315 961',
            'Email: ingenieria.teslaelectricidad@gmail.com',
            'Jr. Las 츼gatas Mz B Lote 09, Urb. San Carlos, SJL'
        ]
        
        for linea in contacto:
            p = self.doc.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(107, 114, 128)
    
    def generar(self, ruta_salida):
        """
        M칠todo abstracto - debe ser implementado por clases hijas
        """
        raise NotImplementedError("Subclases deben implementar generar()")
