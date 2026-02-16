# -*- coding: utf-8 -*-
"""
Generador de Proyecto Complejo PMI - COMPLETO
Project Charter según metodología PMI con KPIs, Gantt, RACI, Stakeholders
100% fidelidad al HTML profesional aprobado
"""

from .base_generator import BaseDocumentGenerator
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ProyectoComplejoPMIGenerator(BaseDocumentGenerator):
    """Generador para proyectos complejos con metodología PMI"""
    
    def _agregar_titulo(self):
        """Agrega título del documento"""
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('PROJECT CHARTER')
        run_titulo.font.size = Pt(30)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        p_subtitulo = self.doc.add_paragraph()
        p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_subtitulo = p_subtitulo.add_run('Gestión de Proyectos según Metodología PMI')
        run_subtitulo.font.size = Pt(14)
        run_subtitulo.font.color.rgb = self.COLOR_SECUNDARIO
        run_subtitulo.font.italic = True
        
        # ✅ Título automático: SERVICIO - INDUSTRIA
        # Buscar primero en opciones, luego en datos
        servicio = self.opciones.get('servicio') or self.datos.get('servicio', 'ELECTRICIDAD')
        industria = self.opciones.get('industria') or self.datos.get('industria', 'CONSTRUCCIÓN')
        
        # Convertir a mayúsculas
        servicio_upper = servicio.upper() if servicio else 'ELECTRICIDAD'
        industria_upper = industria.upper() if industria else 'CONSTRUCCIÓN'
        titulo_proyecto = f"{servicio_upper} - {industria_upper}"
        
        p_nombre = self.doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nombre = p_nombre.add_run(titulo_proyecto)
        run_nombre.font.size = Pt(18)
        run_nombre.font.color.rgb = self.COLOR_SECUNDARIO
        
        codigo = self.datos.get('codigo_proyecto', 'PROY-PMI-001')
        p_codigo = self.doc.add_paragraph()
        p_codigo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_codigo = p_codigo.add_run(codigo)
        run_codigo.font.size = Pt(14)
        run_codigo.font.color.rgb = self.COLOR_SECUNDARIO
        run_codigo.font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_info_grid(self):
        """Agrega grid de información (4 cards)"""
        # Extraer nombre del cliente si es un diccionario
        cliente_data = self.datos.get('cliente', '')
        if isinstance(cliente_data, dict):
            cliente = cliente_data.get('nombre', '')
        else:
            cliente = str(cliente_data) if cliente_data else ''
        duracion = self.datos.get('duracion_total', 60)
        fecha_inicio = self.datos.get('fecha_inicio', '01/01/2025')
        fecha_fin = self.datos.get('fecha_fin', '28/02/2025')
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        cards = [
            ('Duración Total', f'{duracion} días'),
            ('Inicio', fecha_inicio),
            ('Fin Estimado', fecha_fin)
        ]
        
        for idx, (label, value) in enumerate(cards):
            cell = table.rows[0].cells[idx]
            cell.text = ''
            
            p_label = cell.add_paragraph()
            run_label = p_label.add_run(label.upper())
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = RGBColor(107, 114, 128)
            run_label.font.bold = True
            
            p_value = cell.add_paragraph()
            run_value = p_value.add_run(str(value))
            run_value.font.size = Pt(14)
            run_value.font.color.rgb = self.COLOR_PRIMARIO
            run_value.font.bold = True
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _agregar_presupuesto(self):
        """Agrega presupuesto destacado"""
        presupuesto = self.datos.get('presupuesto', '75,000')
        
        # ✅ Extraer moneda de OPCIONES (no de datos)
        moneda = self.opciones.get('moneda', 'USD')
        
        # Determinar símbolo de moneda
        simbolo_moneda = 'S/ ' if moneda == 'PEN' else '€ ' if moneda == 'EUR' else '$ '
        
        p_label = self.doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.add_run('PRESUPUESTO TOTAL DEL PROYECTO')
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = RGBColor(107, 114, 128)
        
        p_valor = self.doc.add_paragraph()
        p_valor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_valor = p_valor.add_run(f'{simbolo_moneda}{presupuesto}')  # ✅ Usar moneda correcta
        run_valor.font.size = Pt(36)
        run_valor.font.color.rgb = self.COLOR_PRIMARIO
        run_valor.font.bold = True
        
        self.doc.add_paragraph()
    
    def _agregar_kpis(self):
        """Agrega KPIs PMI (5 indicadores)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('INDICADORES DE DESEMPEÑO (KPIs)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # KPIs con valores por defecto robustos
        # Intentar obtener de 'kpis' anidado o nivel superior, con fallback a valores seguros
        kpis_data = self.datos.get('kpis', {})
        # Asegurar que kpis_data sea dict
        if not isinstance(kpis_data, dict):
            kpis_data = {}
        
        # Prioridad: kpis anidados > kpis nivel superior > valores por defecto
        # ✅ CORREGIDO: Convertir explícitamente a string para evitar errores de tipo
        spi = str(kpis_data.get('spi') or self.datos.get('spi') or '1.0')
        cpi = str(kpis_data.get('cpi') or self.datos.get('cpi') or '1.0')
        ev_k = str(kpis_data.get('ev_k') or self.datos.get('ev_k') or '0')
        pv_k = str(kpis_data.get('pv_k') or self.datos.get('pv_k') or '0')
        ac_k = str(kpis_data.get('ac_k') or self.datos.get('ac_k') or '0')
        
        # ✅ Extraer moneda de OPCIONES o DATOS con fallback
        moneda = self.opciones.get('moneda') or self.datos.get('moneda') or 'USD'
        simbolo_moneda = 'S/' if moneda == 'PEN' else '€' if moneda == 'EUR' else '$'
        
        table = self.doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'
        
        kpis = [
            ('SPI', spi, 'Schedule Performance'),
            ('CPI', cpi, 'Cost Performance'),
            ('EV', f'{simbolo_moneda}{ev_k}K', 'Earned Value'),  # ✅ Usar moneda dinámica
            ('PV', f'{simbolo_moneda}{pv_k}K', 'Planned Value'),  # ✅ Usar moneda dinámica
            ('AC', f'{simbolo_moneda}{ac_k}K', 'Actual Cost')  # ✅ Usar moneda dinámica
        ]
        
        for idx, (label, value, desc) in enumerate(kpis):
            cell_label = table.rows[0].cells[idx]
            cell_label.text = label
            cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_label.paragraphs[0].runs[0].font.size = Pt(12)
            cell_label.paragraphs[0].runs[0].font.bold = True
            
            cell_value = table.rows[1].cells[idx]
            cell_value.text = value
            cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_value.paragraphs[0].runs[0].font.size = Pt(20)
            cell_value.paragraphs[0].runs[0].font.color.rgb = self.COLOR_PRIMARIO
            cell_value.paragraphs[0].runs[0].font.bold = True
            
            cell_desc = table.rows[2].cells[idx]
            cell_desc.text = desc
            cell_desc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_desc.paragraphs[0].runs[0].font.size = Pt(9)
            cell_desc.paragraphs[0].runs[0].font.color.rgb = RGBColor(156, 163, 175)
        
        self.doc.add_paragraph()
    
    def _agregar_alcance(self):
        """Agrega alcance del proyecto (WBS Level 1)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ALCANCE DEL PROYECTO (WBS Level 1)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        alcance = self.datos.get('alcance_proyecto', 'Descripción del alcance del proyecto...')
        p_alcance = self.doc.add_paragraph(alcance)
        p_alcance.runs[0].font.size = Pt(12)
        p_alcance.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
    
    def _agregar_cronograma_gantt(self):
        """Agrega cronograma Gantt con barras visuales"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('CRONOGRAMA DEL PROYECTO (Diagrama Gantt)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        self.doc.add_paragraph()  # Espacio
        
        # ✅ USAR DATOS DINÁMICOS
        fases = self.datos.get('cronograma_fases', [])
        duracion_total = self.datos.get('duracion_total', 0)
        
        if not fases:
            # Fallback por si acaso
            dias_ing = self.datos.get('dias_ingenieria', 15)
            dias_ejec = self.datos.get('dias_ejecucion', 25)
            fases = [
                {'label': '1. Inicio y Planificación', 'dias': '10 días', 'width': '15%'},
                {'label': '2. Gestión Stakeholders', 'dias': '3 días', 'width': '8%'},
                {'label': '3. Ingeniería y Diseño', 'dias': f'{dias_ing} días', 'width': '25%'},
                {'label': '4. Ejecución', 'dias': f'{dias_ejec} días', 'width': '40%'},
                {'label': '5. Pruebas y Puesta en Marcha', 'dias': '8 días', 'width': '10%'},
                {'label': '6. Cierre', 'dias': '5 días', 'width': '7%'}
            ]
            duracion_total = 10 + 3 + dias_ing + dias_ejec + 8 + 5
        
        # Calcular duración total si no está disponible
        if duracion_total == 0:
            for fase in fases:
                dias_str = fase.get('dias', '0')
                dias_num = int(''.join(filter(str.isdigit, str(dias_str)))) if dias_str else 0
                duracion_total += dias_num
        
        # ✅ CREAR TABLA CON BARRAS VISUALES
        # Tabla de 3 columnas: Nombre | Barra Visual | Días
        table = self.doc.add_table(rows=len(fases), cols=3)
        table.style = 'Light Grid Accent 1'
        
        for i, item in enumerate(fases):
            # Manejar tanto formato diccionario como tupla
            if isinstance(item, dict):
                label = item.get('label', item.get('nombre', ''))
                dias = item.get('dias', '')
                width_percent = item.get('width', '10%')
            elif isinstance(item, (list, tuple)) and len(item) >= 2:
                label, dias = item[0], item[1]
                width_percent = '10%'
            else:
                continue
            
            # Extraer número de días
            dias_num = int(''.join(filter(str.isdigit, str(dias)))) if dias else 0
            
            # Calcular porcentaje de ancho de la barra
            if duracion_total > 0:
                bar_width_percent = int((dias_num / duracion_total) * 100)
            else:
                bar_width_percent = 10
            
            # Fila de la tabla
            row = table.rows[i]
            
            # Columna 1: Nombre de la fase
            cell_nombre = row.cells[0]
            cell_nombre.width = Inches(2.5)
            p_nombre = cell_nombre.paragraphs[0]
            run_nombre = p_nombre.add_run(label)
            run_nombre.font.size = Pt(10)
            run_nombre.font.bold = False
            run_nombre.font.color.rgb = RGBColor(55, 65, 81)
            
            # Columna 2: Barra visual
            cell_barra = row.cells[1]
            cell_barra.width = Inches(3.5)
            
            # Crear tabla interna para la barra (10 celdas = 100%)
            inner_table = cell_barra.add_table(rows=1, cols=10)
            inner_table.autofit = False
            
            # Calcular cuántas celdas colorear
            cells_to_color = max(1, min(10, int(bar_width_percent / 10)))
            
            for j in range(10):
                inner_cell = inner_table.rows[0].cells[j]
                inner_cell.width = Inches(0.35)
                
                # Colorear celdas según el porcentaje
                if j < cells_to_color:
                    # Aplicar color de fondo usando el color primario
                    shading_elm = OxmlElement('w:shd')
                    color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
                    shading_elm.set(qn('w:fill'), color_hex)
                    inner_cell._element.get_or_add_tcPr().append(shading_elm)
                else:
                    # Celdas vacías con patrón de líneas diagonales semi-transparentes
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F9FAFB')  # Gris muy claro
                    shading_elm.set(qn('w:val'), 'thinDiagStripe')  # Patrón de líneas diagonales
                    shading_elm.set(qn('w:color'), 'E5E7EB')  # Color de las líneas (gris medio)
                    inner_cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Ajustar altura de celda
                inner_cell.height = Inches(0.25)
            
            # Columna 3: Días
            cell_dias = row.cells[2]
            cell_dias.width = Inches(0.8)
            p_dias = cell_dias.paragraphs[0]
            run_dias = p_dias.add_run(dias)
            run_dias.font.size = Pt(11)
            run_dias.font.bold = True
            run_dias.font.color.rgb = self.COLOR_PRIMARIO
        
        # ✅ Agregar duración total estimada
        if duracion_total > 0:
            self.doc.add_paragraph()  # Espacio
            p_total = self.doc.add_paragraph()
            run_total = p_total.add_run(f'DURACIÓN TOTAL ESTIMADA: {duracion_total} días hábiles')
            run_total.font.size = Pt(12)
            run_total.font.bold = True
            run_total.font.color.rgb = self.COLOR_ACENTO
        
        self.doc.add_paragraph()

        self.doc.add_paragraph()

    def _agregar_stakeholders(self):
        """Agrega registro de stakeholders"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('REGISTRO DE STAKEHOLDERS')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # ✅ USAR DATOS DINÁMICOS
        stakeholders = self.datos.get('stakeholders', [])
        
        if not stakeholders:
             stakeholders = [
                {'nombre': 'Cliente', 'rol': 'Cliente / Patrocinador Principal', 'poder': 'Alto', 'interes': 'Alto'},
                {'nombre': 'Jefe de Proyecto', 'rol': 'Project Manager', 'poder': 'Alto', 'interes': 'Alto'},
                {'nombre': 'Equipo Técnico', 'rol': 'Ingenieros y Técnicos', 'poder': 'Medio', 'interes': 'Alto'}
            ]
        
        for sh in stakeholders:
            # Manejar si es string o dict
            if isinstance(sh, dict):
                nombre = sh.get('nombre', '')
                rol = sh.get('rol', '')
                poder = sh.get('poder', '')
                interes = sh.get('interes', '')
            else:
                nombre = str(sh)
                rol = ''
                poder = ''
                interes = ''

            p_sh = self.doc.add_paragraph()
            run_nombre = p_sh.add_run(f"{nombre}: ")
            run_nombre.font.size = Pt(14)
            run_nombre.font.bold = True
            run_nombre.font.color.rgb = self.COLOR_PRIMARIO
            
            run_rol = p_sh.add_run(rol)
            run_rol.font.size = Pt(11)
            
            if poder or interes:
                p_badges = self.doc.add_paragraph()
                run_badges = p_badges.add_run(f"Poder: {poder} | Interés: {interes}")
                run_badges.font.size = Pt(10)
                run_badges.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()

    def _agregar_matriz_raci(self):
        """Agrega matriz RACI"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('MATRIZ RACI (Responsabilidades)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # ✅ USAR DATOS DINÁMICOS
        actividades = self.datos.get('raci_actividades', [])
        
        if not actividades:
            # Fallback
            actividades = [
                {'actividad': 'Planificación del Proyecto', 'roles': ['A', 'R', 'I', 'C', 'C']},
                {'actividad': 'Diseño e Ingeniería', 'roles': ['A', 'R', 'C', 'C', 'I']},
                {'actividad': 'Ejecución de Obra', 'roles': ['A', 'A', 'R', 'C', 'I']},
                {'actividad': 'Control de Calidad', 'roles': ['A', 'C', 'C', 'R', 'I']},
                {'actividad': 'Aprobación de Entregables', 'roles': ['R', 'C', 'I', 'C', 'A']}
            ]

        # Validar estructura y obtener max columnas
        max_roles = 5
        if actividades and isinstance(actividades[0], dict) and 'roles' in actividades[0]:
            max_roles = max(len(a['roles']) for a in actividades)
        
        table = self.doc.add_table(rows=len(actividades) + 1, cols=max_roles + 1)
        table.style = 'Table Grid'
        
        # Header
        headers = ['Actividad', 'PM', 'Ing. Residente', 'Técnicos', 'Inspector QA', 'Cliente']
        # Ajustar headers si hay más o menos columnas
        if len(headers) < max_roles + 1:
            headers.extend([f'Rol {i+1}' for i in range(len(headers)-1, max_roles)])
            
        for idx in range(max_roles + 1):
            cell = table.rows[0].cells[idx]
            header_text = headers[idx] if idx < len(headers) else f'Rol {idx}'
            cell.text = header_text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        for idx, item in enumerate(actividades, 1):
            if isinstance(item, dict):
                nombre = item.get('actividad', '')
                roles = item.get('roles', [])
            elif isinstance(item, (list, tuple)):
                nombre = item[0]
                roles = item[1]
            else:
                continue
                
            table.rows[idx].cells[0].text = nombre
            for j, rol in enumerate(roles):
                if j + 1 < len(table.rows[idx].cells):
                    cell = table.rows[idx].cells[j+1]
                    cell.text = str(rol)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Leyenda
        p_leyenda = self.doc.add_paragraph()
        run_leyenda = p_leyenda.add_run('Leyenda: R=Responsable | A=Aprobador | C=Consultado | I=Informado')
        run_leyenda.font.size = Pt(10)
        
        self.doc.add_paragraph()

        self.doc.add_paragraph()

    def _agregar_riesgos(self):
        """Agrega registro de riesgos (Top 5)"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('REGISTRO DE RIESGOS (Top 5)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # ✅ USAR DATOS DINÁMICOS
        riesgos = self.datos.get('riesgos', [])
        
        if not riesgos:
            riesgos = [
                {'id': 'R01', 'descripcion': 'Retrasos en entrega de equipos', 'probabilidad': 'Media', 'impacto': 'Alto', 'severidad': 'Alta', 'mitigacion': 'Compra anticipada'},
                {'id': 'R02', 'descripcion': 'Variación de precios', 'probabilidad': 'Media', 'impacto': 'Medio', 'severidad': 'Media', 'mitigacion': 'Precio fijo'},
                {'id': 'R03', 'descripcion': 'Cambios en alcance', 'probabilidad': 'Alta', 'impacto': 'Alto', 'severidad': 'Alta', 'mitigacion': 'Control de cambios'}
            ]
        
        table = self.doc.add_table(rows=len(riesgos) + 1, cols=6)
        table.style = 'Table Grid'
        
        # Header
        headers = ['ID', 'Riesgo', 'Probabilidad', 'Impacto', 'Severidad', 'Plan de Mitigación']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Riesgos
        for idx, riesgo in enumerate(riesgos, 1):
            if isinstance(riesgo, dict):
                 id_r = riesgo.get('id', str(idx))
                 desc = riesgo.get('descripcion', '')
                 prob = riesgo.get('probabilidad', '')
                 imp = riesgo.get('impacto', '')
                 sev = riesgo.get('severidad', '')
                 mit = riesgo.get('mitigacion', '')
            else:
                 id_r = str(idx)
                 desc = str(riesgo)
                 prob = ''
                 imp = ''
                 sev = ''
                 mit = ''

            row = table.rows[idx]
            row.cells[0].text = id_r
            row.cells[1].text = desc
            row.cells[2].text = prob
            row.cells[3].text = imp
            row.cells[4].text = sev
            row.cells[5].text = mit
            
            for j in [2, 3, 4]:
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()

    def _agregar_entregables(self):
        """Agrega entregables principales"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('ENTREGABLES PRINCIPALES DEL PROYECTO')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # ✅ USAR DATOS DINÁMICOS
        entregables = self.datos.get('entregables', [])
        
        if not entregables:
             entregables = [
                '📋 Project Charter', '📊 Plan de gestión', '👥 Registro stakeholders',
                '📝 WBS y diccionario', '📅 Cronograma Gantt', '📐 Planos instalación',
                '📄 Especif. técnicas', '✅ Plan de calidad', '⚠️ Registro de riesgos',
                '🔬 Protocolos FAT/SAT', '🏗️ Planos as-built', '📚 Lecciones aprendidas',
                '🎯 Acta de cierre'
            ]
        
        for entregable in entregables:
            p_ent = self.doc.add_paragraph(f'• {entregable}')
            p_ent.runs[0].font.size = Pt(11)
        
        self.doc.add_paragraph()

    def _agregar_normativa(self):
        """Agrega normativa aplicable"""
        normativa = self.datos.get('normativa_aplicable', 'CNE - Código Nacional de Electricidad')
        subtitulo = self.datos.get('subtitulo_normativa', 'Gestión según PMBOK® Guide 7th Edition')
        
        p_label = self.doc.add_paragraph()
        run_label = p_label.add_run('NORMATIVA Y ESTÁNDARES APLICABLES')
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = RGBColor(107, 114, 128)
        
        p_valor = self.doc.add_paragraph()
        run_valor = p_valor.add_run(normativa)
        run_valor.font.size = Pt(14)
        run_valor.font.color.rgb = self.COLOR_PRIMARIO
        run_valor.font.bold = True
        
        p_pmbok = self.doc.add_paragraph()
        run_pmbok = p_pmbok.add_run(subtitulo)
        run_pmbok.font.size = Pt(11)
        run_pmbok.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
    def _agregar_matriz_raci(self):
        """Agrega Matriz RACI si la complejidad lo requiere"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('MATRIZ DE ASIGNACIÓN DE RESPONSABILIDADES (RACI)')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        # Datos
        raci_data = self.datos.get('raci_actividades', [])
        if not raci_data:
            raci_data = [
                 {'actividad': 'Planificación', 'roles': ['A', 'R', 'I', 'C', 'C']},
                 {'actividad': 'Ejecución', 'roles': ['A', 'A', 'R', 'C', 'I']},
                 {'actividad': 'Cierre', 'roles': ['R', 'C', 'I', 'C', 'A']}
            ]
            
        # Tabla
        table = self.doc.add_table(rows=len(raci_data) + 1, cols=6)
        table.style = 'Table Grid'
        
        # Header
        headers = ['Actividad', 'PM', 'Residente', 'Técnicos', 'QA', 'Cliente']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = OxmlElement('w:shd')
            color_hex = self._rgb_to_hex(self.COLOR_PRIMARIO)
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
            
        # Filas
        for idx, row_data in enumerate(raci_data, 1):
            actividad = row_data.get('actividad', '')
            roles = row_data.get('roles', [])
            
            row = table.rows[idx]
            row.cells[0].text = actividad
            
            for j in range(5):
                if j < len(roles):
                    cell = row.cells[j+1]
                    cell.text = roles[j]
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        self.doc.add_paragraph()

    def _agregar_recursos(self):
        """Agrega sección de Recursos y Materiales"""
        p_titulo = self.doc.add_paragraph()
        run_titulo = p_titulo.add_run('RECURSOS Y MATERIALES CRÍTICOS')
        run_titulo.font.size = Pt(18)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = self.COLOR_PRIMARIO
        
        recursos = self.datos.get('recursos_humanos', [])
        materiales = self.datos.get('materiales', [])
        
        # 1. Recursos Humanos
        p_sub1 = self.doc.add_paragraph()
        run_sub1 = p_sub1.add_run('Equipo del Proyecto')
        run_sub1.font.bold = True
        run_sub1.font.color.rgb = self.COLOR_SECUNDARIO
        run_sub1.font.size = Pt(14)
        
        if not recursos:
            recursos = ['Project Manager', 'Ingeniero Residente', 'Técnicos']
            
        for rec in recursos:
            self.doc.add_paragraph(f'• {rec}', style='List Bullet')
            
        self.doc.add_paragraph()
        
        # 2. Materiales
        p_sub2 = self.doc.add_paragraph()
        run_sub2 = p_sub2.add_run('Materiales Principales')
        run_sub2.font.bold = True
        run_sub2.font.color.rgb = self.COLOR_SECUNDARIO
        run_sub2.font.size = Pt(14)
        
        if not materiales:
            materiales = ['Equipos principales', 'Materiales de instalación']
            
        for mat in materiales:
            self.doc.add_paragraph(f'• {mat}', style='List Bullet')
            
        self.doc.add_paragraph()
    def generar(self, ruta_salida):
        """Genera el documento completo con proteccion de fallos"""
        import logging
        logger = logging.getLogger(__name__)
        logger.info("🚀 INICIANDO GENERACIÓN PMI (MODO DEFENSIVO)...")
        
        # ✅ NUEVO: Detectar complejidad del proyecto
        complejidad = self.datos.get('complejidad', 7)
        logger.info(f"📊 Complejidad del proyecto: {complejidad} fases")
        
        # Secciones comunes (siempre se generan)
        methods_comunes = [
            (self._agregar_header_basico, "Header Básico"),
            (self._agregar_titulo, "Título"),
            (self._agregar_info_grid, "Info Grid"),
            (self._agregar_presupuesto, "Presupuesto"),
            (self._agregar_kpis, "KPIs"),
            (self._agregar_alcance, "Alcance"),
            (self._agregar_cronograma_gantt, "Gantt"),
            (self._agregar_stakeholders, "Stakeholders"),
        ]
        
        # ✅ Matriz RACI: Solo si complejidad >= 6
        if complejidad >= 6:
            logger.info("✅ Incluyendo Matriz RACI (complejidad >= 6)")
            methods_comunes.append((self._agregar_matriz_raci, "RACI"))
        else:
            logger.info("⏭️ Omitiendo Matriz RACI (complejidad < 6)")
        
        # Continuar con secciones comunes
        methods_comunes.extend([
            (self._agregar_riesgos, "Riesgos"),
            (self._agregar_recursos, "Recursos"),  # ✅ NUEVO
            (self._agregar_entregables, "Entregables"),
            (self._agregar_normativa, "Normativa"),
            (self._agregar_footer_basico, "Footer")
        ])

        for method, name in methods_comunes:
            try:
                logger.info(f"👉 Ejecutando: {name}")
                method()
            except Exception as e:
                logger.error(f"❌ ERROR EN {name.upper()}: {e}", exc_info=True)
                # No relanzamos para que se genere al menos parcial
                p_error = self.doc.add_paragraph()
                run_error = p_error.add_run(f"[ERROR GENERANDO SECCIÓN {name}: {e}]")
                run_error.font.color.rgb = RGBColor(255, 0, 0)
                run_error.font.bold = True
        
        try:
            logger.info(f"💾 Guardando documento en: {ruta_salida}")
            self.doc.save(str(ruta_salida))
            return ruta_salida
        except Exception as e:
            logger.error(f"❌ ERROR FINAL GUARDANDO ARCHIVO: {e}", exc_info=True)
            raise e


def generar_proyecto_complejo_pmi(datos, ruta_salida, opciones=None):
    """Función de entrada para generar proyecto complejo PMI"""
    generator = ProyectoComplejoPMIGenerator(datos, opciones)
    return generator.generar(ruta_salida)
