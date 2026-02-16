"""
Generation Router - API Endpoints for Document Generation
Handles PDF and Word export functionality
"""
from fastapi import APIRouter, HTTPException, Response
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import Optional, Dict, Any
import io
import os
from datetime import datetime

# Import generators
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    print("WeasyPrint not available (GTK3 missing?). Will try ReportLab.")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Fallback PDF generator
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

from utils.template_generator import generate_html_template

router = APIRouter(prefix="/api/generate", tags=["generation"])

class GenerateRequest(BaseModel):
    title: str
    type: str
    data: Dict[str, Any]
    color_scheme: Optional[str] = 'azul-tesla'
    font: Optional[str] = 'Calibri'
    user_id: Optional[str] = None

    class Config:
        extra = "ignore"

@router.post("/pdf")
async def generate_pdf(request: GenerateRequest):
    """Generate PDF document from data"""
    try:
        pdf_file = io.BytesIO()
        generated = False
        
        # Strategy 1: WeasyPrint (High Quality)
        if WEASYPRINT_AVAILABLE:
            try:
                html_content = generate_html_template(
                    request.type, 
                    request.data, 
                    request.color_scheme,
                    request.font
                )
                HTML(string=html_content).write_pdf(pdf_file)
                pdf_file.seek(0)
                generated = True
            except Exception as wp_error:
                print(f"WeasyPrint failed: {wp_error}. Falling back to ReportLab...")
                generated = False
                
        # Strategy 2: ReportLab (Reliable Fallback)
        if not generated and REPORTLAB_AVAILABLE:
            c = canvas.Canvas(pdf_file, pagesize=A4)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 800, request.title or "Documento Generado")
            
            c.setFont("Helvetica", 12)
            y = 750
            c.drawString(50, y, f"Tipo: {request.type}")
            y -= 20
            c.drawString(50, y, f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
            y -= 40
            
            # Simple content dump
            c.drawString(50, y, "Datos del Documento:")
            y -= 25
            
            # Iterate safely through basic data
            text_object = c.beginText(50, y)
            text_object.setFont("Helvetica", 10)
            
            if request.data.get('cliente'):
                text_object.textLine("CLIENTE:")
                client_data = request.data['cliente']
                if isinstance(client_data, dict):
                    for k, v in client_data.items():
                        text_object.textLine(f" - {k}: {v}")
                else:
                    text_object.textLine(f" - {str(client_data)}")
                text_object.textLine(" ")

            if request.data.get('proyecto'):
                text_object.textLine("PROYECTO:")
                project_data = request.data['proyecto']
                if isinstance(project_data, dict):
                    for k, v in project_data.items():
                        text_object.textLine(f" - {k}: {v}")
                else:
                     text_object.textLine(f" - {str(project_data)}")
            
            c.drawText(text_object)
            c.showPage()
            c.save()
            pdf_file.seek(0)
            generated = True
            
        if not generated:
             raise HTTPException(
                status_code=500, 
                detail="No PDF generation libraries available (or both failed)"
            )

        # 3. Return as downloadable file
        filename = f"{request.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
        return Response(
            content=pdf_file.getvalue(), 
            headers=headers, 
            media_type="application/pdf"
        )
        
    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating PDF: {str(e)}")

@router.post("/word")
async def generate_word(request: GenerateRequest):
    """Generate Word document from data"""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500, 
            detail="Word generation library (python-docx) not installed"
        )
        
    try:
        # Create Word document
        doc = Document()
        
        # Add basic content (simplified for now, mimicking structure)
        doc.add_heading(request.title, 0)
        
        p = doc.add_paragraph()
        p.add_run(f"Tipo: {request.type}\n").bold = True
        p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")
        
        # Add data dump for now (robust logic would go in a utility)
        doc.add_heading('Datos del Documento', level=1)
        
        # Client info
        if 'cliente' in request.data:
            doc.add_heading('Cliente', level=2)
            client_data = request.data['cliente']
            if isinstance(client_data, dict):
                for key, value in client_data.items():
                    doc.add_paragraph(f"{key.capitalize()}: {value}")
            else:
                 doc.add_paragraph(str(client_data))
                
        # Project info
        if 'proyecto' in request.data:
            doc.add_heading('Proyecto', level=2)
            project_data = request.data['proyecto']
            if isinstance(project_data, dict):
                for key, value in project_data.items():
                    doc.add_paragraph(f"{key.capitalize()}: {value}")
            else:
                doc.add_paragraph(str(project_data))
        
        # Save to buffer
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        filename = f"{request.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
        return Response(
            content=docx_file.getvalue(),
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print(f"Error generating Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating Word doc: {str(e)}")

@router.post("/excel")
async def generate_excel(request: GenerateRequest):
    """Generate Excel spreadsheet from data"""
    if not OPENPYXL_AVAILABLE:
        raise HTTPException(
            status_code=500, 
            detail="Excel generation library (openpyxl) not installed"
        )
        
    try:
        wb = Workbook()
        
        # --- SHEET 1: RESUMEN ---
        ws = wb.active
        ws.title = "Resumen del Proyecto"
        
        # Styles
        title_font = Font(name='Calibri', size=14, bold=True, color='FFFFFF')
        header_font = Font(name='Calibri', size=11, bold=True)
        fill_blue = PatternFill(start_color='1E40AF', end_color='1E40AF', fill_type='solid')
        fill_gray = PatternFill(start_color='F3F4F6', end_color='F3F4F6', fill_type='solid')
        border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        # Title
        ws['A1'] = request.title.upper()
        ws.merge_cells('A1:E1')
        ws['A1'].font = title_font
        ws['A1'].fill = fill_blue
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # Info Block
        current_row = 3
        ws[f'A{current_row}'] = "INFORMACIÓN GENERAL"
        ws[f'A{current_row}'].font = header_font
        current_row += 1
        
        ws[f'A{current_row}'] = "Tipo de Documento:"
        ws[f'B{current_row}'] = request.type.replace('-', ' ').title()
        current_row += 1
        
        ws[f'A{current_row}'] = "Fecha de Emisión:"
        ws[f'B{current_row}'] = datetime.now().strftime('%d/%m/%Y')
        current_row += 2
        
        # Client Data
        if request.data.get('cliente'):
            ws[f'A{current_row}'] = "DATOS DEL CLIENTE"
            ws[f'A{current_row}'].font = header_font
            current_row += 1
            
            client_data = request.data['cliente']
            if isinstance(client_data, dict):
                for k, v in client_data.items():
                    ws[f'A{current_row}'] = str(k).replace('_', ' ').title()
                    ws[f'B{current_row}'] = str(v)
                    current_row += 1
            else:
                ws[f'A{current_row}'] = "Cliente"
                ws[f'B{current_row}'] = str(client_data)
                current_row += 1
            current_row += 1

        # Project Data
        if request.data.get('proyecto'):
            ws[f'A{current_row}'] = "DATOS DEL PROYECTO"
            ws[f'A{current_row}'].font = header_font
            current_row += 1
            
            project_data = request.data['proyecto']
            if isinstance(project_data, dict):
                for k, v in project_data.items():
                    if isinstance(v, (str, int, float)):
                        ws[f'A{current_row}'] = str(k).replace('_', ' ').title()
                        ws[f'B{current_row}'] = str(v)
                        current_row += 1
            else:
                # It's a string
                ws[f'A{current_row}'] = "Nombre del Proyecto"
                ws[f'B{current_row}'] = str(project_data)
                current_row += 1
        
        # Adjust Columns
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 40

        # --- SHEET 2: DETALLE ---
        if request.data.get('items'):
            ws2 = wb.create_sheet("Detalle de Items")
            
            headers = ["Item", "Descripción", "Unidad", "Cantidad", "Precio U.", "Total"]
            for col, header in enumerate(headers, 1):
                cell = ws2.cell(row=1, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = fill_gray
                cell.border = border
            
            row_idx = 2
            total_general = 0
            
            for item in request.data['items']:
                # Ensure data types
                desc = item.get('descripcion', 'Item')
                unidad = item.get('unidad', 'und')
                cant = float(item.get('cantidad', 1))
                precio = float(item.get('precio_unitario', 0))
                total = cant * precio
                total_general += total
                
                ws2.cell(row=row_idx, column=1, value=row_idx-1).border = border
                ws2.cell(row=row_idx, column=2, value=desc).border = border
                ws2.cell(row=row_idx, column=3, value=unidad).border = border
                ws2.cell(row=row_idx, column=4, value=cant).border = border
                ws2.cell(row=row_idx, column=5, value=precio).border = border
                ws2.cell(row=row_idx, column=6, value=total).border = border
                row_idx += 1
            
            # Total Check
            ws2.cell(row=row_idx, column=5, value="TOTAL:").font = header_font
            ws2.cell(row=row_idx, column=6, value=total_general).font = header_font
            
            # Formats
            for col in ['E', 'F']:
                for cell in ws2[col]:
                    cell.number_format = '#,##0.00'
            
            ws2.column_dimensions['B'].width = 50
        
        # Save
        excel_file = io.BytesIO()
        wb.save(excel_file)
        excel_file.seek(0)
        
        filename = f"{request.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
        return Response(
            content=excel_file.getvalue(),
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except Exception as e:
        print(f"Error generating Excel: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error creating Excel: {str(e)}")
